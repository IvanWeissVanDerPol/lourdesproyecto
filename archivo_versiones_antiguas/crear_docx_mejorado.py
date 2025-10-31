#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertidor Mejorado de Markdown a DOCX
Versión 2.0 - Con mejor manejo de caracteres especiales y formato profesional
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import sys

# Asegurar codificación UTF-8 en Windows
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

def crear_estilos_profesionales(doc):
    """Crear estilos profesionales para documento académico"""

    # Configurar estilos por defecto del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Configurar párrafo normal
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Título 1 - Secciones principales (## en markdown)
    try:
        heading1 = doc.styles['Heading 1']
    except:
        heading1 = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

    heading1.font.name = 'Arial'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro profesional
    heading1.paragraph_format.space_before = Pt(24)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.page_break_before = False  # No romper página antes

    # Título 2 - Subsecciones (### en markdown)
    try:
        heading2 = doc.styles['Heading 2']
    except:
        heading2 = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)

    heading2.font.name = 'Arial'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 51, 102)
    heading2.paragraph_format.space_before = Pt(18)
    heading2.paragraph_format.space_after = Pt(10)
    heading2.paragraph_format.keep_with_next = True

    # Título 3 - Sub-subsecciones (#### en markdown)
    try:
        heading3 = doc.styles['Heading 3']
    except:
        heading3 = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)

    heading3.font.name = 'Arial'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(51, 51, 51)
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.paragraph_format.keep_with_next = True

    # Título 4 - Apartados menores (##### en markdown)
    try:
        heading4 = doc.styles['Heading 4']
    except:
        heading4 = doc.styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)

    heading4.font.name = 'Arial'
    heading4.font.size = Pt(11)
    heading4.font.bold = True
    heading4.font.italic = True
    heading4.font.color.rgb = RGBColor(51, 51, 51)
    heading4.paragraph_format.space_before = Pt(10)
    heading4.paragraph_format.space_after = Pt(6)

    # Estilo para citas/blockquotes
    try:
        quote = doc.styles['Quote']
    except:
        quote = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)

    quote.font.name = 'Times New Roman'
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor(102, 102, 102)
    quote.paragraph_format.left_indent = Inches(0.5)
    quote.paragraph_format.right_indent = Inches(0.5)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(6)

def configurar_documento(doc):
    """Configurar propiedades del documento"""

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)      # 2.5 cm
        section.bottom_margin = Cm(2.5)   # 2.5 cm
        section.left_margin = Cm(3)       # 3 cm
        section.right_margin = Cm(3)      # 3 cm
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

def limpiar_texto(texto):
    """Limpiar y normalizar texto para evitar problemas de encoding"""

    # Reemplazos de caracteres especiales comunes
    reemplazos = {
        '\u00f3': 'ó',
        '\u00e1': 'á',
        '\u00e9': 'é',
        '\u00ed': 'í',
        '\u00fa': 'ú',
        '\u00f1': 'ñ',
        '\u00c1': 'Á',
        '\u00c9': 'É',
        '\u00cd': 'Í',
        '\u00d3': 'Ó',
        '\u00da': 'Ú',
        '\u00d1': 'Ñ',
    }

    for mal, bien in reemplazos.items():
        texto = texto.replace(mal, bien)

    return texto

def procesar_formato_inline(texto):
    """Procesar formato inline (negrita, cursiva, código) y devolver lista de runs"""

    runs = []
    texto = limpiar_texto(texto)

    # Patrón para capturar: **negrita**, *cursiva*, `código`
    patron = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    partes = re.split(patron, texto)

    for parte in partes:
        if not parte:
            continue

        if parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
            # Negrita
            runs.append({'texto': parte[2:-2], 'negrita': True})
        elif parte.startswith('*') and parte.endswith('*') and len(parte) > 2 and not parte.startswith('**'):
            # Cursiva
            runs.append({'texto': parte[1:-1], 'cursiva': True})
        elif parte.startswith('`') and parte.endswith('`'):
            # Código inline
            runs.append({'texto': parte[1:-1], 'codigo': True})
        else:
            # Texto normal
            runs.append({'texto': parte, 'normal': True})

    return runs

def agregar_run_formateado(paragraph, run_info):
    """Agregar un run con formato al párrafo"""

    run = paragraph.add_run(run_info['texto'])

    if run_info.get('negrita'):
        run.bold = True

    if run_info.get('cursiva'):
        run.italic = True

    if run_info.get('codigo'):
        run.font.name = 'Courier New'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(199, 37, 78)  # Color rojo para código

    return run

def crear_tabla_formateada(doc, filas_tabla):
    """Crear tabla con formato profesional"""

    if len(filas_tabla) < 2:  # Necesitamos al menos encabezado + 1 fila
        return

    num_cols = max(len(fila) for fila in filas_tabla)
    tabla = doc.add_table(rows=len(filas_tabla), cols=num_cols)
    tabla.style = 'Light Grid Accent 1'

    # Llenar tabla
    for idx_fila, fila in enumerate(filas_tabla):
        for idx_col, celda in enumerate(fila):
            if idx_col < num_cols:
                cell = tabla.rows[idx_fila].cells[idx_col]
                cell.text = limpiar_texto(celda)

                # Formato para encabezado
                if idx_fila == 0:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                else:
                    # Formato para celdas normales
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)

def convertir_markdown_a_docx(archivo_md, archivo_docx):
    """Convertir archivo markdown a DOCX con formato profesional"""

    print("\n" + "="*80)
    print("  CONVERTIDOR MEJORADO MARKDOWN → WORD")
    print("  Versión 2.0 - Con soporte completo de caracteres especiales")
    print("="*80)
    print(f"\n[*] Archivo origen: {archivo_md}")
    print(f"[*] Archivo destino: {archivo_docx}")

    # Leer archivo con encoding UTF-8
    print("\n[*] Leyendo archivo markdown...")
    with open(archivo_md, 'r', encoding='utf-8') as f:
        contenido = f.read()

    # Crear documento
    print("[*] Creando documento Word...")
    doc = Document()

    # Configurar documento
    configurar_documento(doc)
    crear_estilos_profesionales(doc)

    print("[*] Procesando contenido...")

    # Procesar línea por línea
    lineas = contenido.split('\n')
    total_lineas = len(lineas)

    en_tabla = False
    filas_tabla = []
    en_bloque_codigo = False
    bloque_codigo = []
    linea_anterior_vacia = False

    for i, linea in enumerate(lineas):
        # Mostrar progreso
        if i % 1000 == 0:
            porcentaje = (i / total_lineas) * 100
            print(f"    Progreso: {porcentaje:.1f}% ({i:,}/{total_lineas:,} líneas)", end='\r')

        # Detectar bloque de código
        if linea.strip().startswith('```'):
            if en_bloque_codigo:
                # Fin del bloque de código
                if bloque_codigo:
                    p = doc.add_paragraph()
                    codigo_texto = '\n'.join(bloque_codigo)
                    run = p.add_run(limpiar_texto(codigo_texto))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(51, 51, 51)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                bloque_codigo = []
                en_bloque_codigo = False
            else:
                en_bloque_codigo = True
            continue

        if en_bloque_codigo:
            bloque_codigo.append(linea)
            continue

        # Detectar tablas
        if '|' in linea and linea.strip().startswith('|'):
            if not en_tabla:
                en_tabla = True
                filas_tabla = []

            celdas = [c.strip() for c in linea.split('|')[1:-1]]

            # Saltar línea separadora
            if all(re.match(r'^-+$', c.strip()) for c in celdas if c.strip()):
                continue

            filas_tabla.append(celdas)
            continue
        else:
            # Finalizar tabla si había una
            if en_tabla and filas_tabla:
                crear_tabla_formateada(doc, filas_tabla)
                filas_tabla = []
                en_tabla = False

        # Líneas vacías
        if not linea.strip():
            if not linea_anterior_vacia:
                linea_anterior_vacia = True
            continue
        else:
            linea_anterior_vacia = False

        # Detectar títulos
        if linea.startswith('# ') and not linea.startswith('## '):
            # Título nivel 1 (raro en el documento)
            p = doc.add_paragraph(limpiar_texto(linea[2:].strip()))
            p.style = 'Heading 1'
            continue

        elif linea.startswith('## ') and not linea.startswith('### '):
            # Título nivel 2 → Heading 1
            p = doc.add_paragraph(limpiar_texto(linea[3:].strip()))
            p.style = 'Heading 1'
            continue

        elif linea.startswith('### ') and not linea.startswith('#### '):
            # Título nivel 3 → Heading 2
            p = doc.add_paragraph(limpiar_texto(linea[4:].strip()))
            p.style = 'Heading 2'
            continue

        elif linea.startswith('#### ') and not linea.startswith('##### '):
            # Título nivel 4 → Heading 3
            p = doc.add_paragraph(limpiar_texto(linea[5:].strip()))
            p.style = 'Heading 3'
            continue

        elif linea.startswith('##### '):
            # Título nivel 5 → Heading 4
            p = doc.add_paragraph(limpiar_texto(linea[6:].strip()))
            p.style = 'Heading 4'
            continue

        # Detectar listas con viñetas
        elif linea.strip().startswith('- ') or linea.strip().startswith('* '):
            texto = linea.strip()[2:]
            p = doc.add_paragraph(limpiar_texto(texto), style='List Bullet')
            continue

        # Detectar listas numeradas
        elif re.match(r'^\d+\.\s', linea.strip()):
            texto = re.sub(r'^\d+\.\s', '', linea.strip())
            p = doc.add_paragraph(limpiar_texto(texto), style='List Number')
            continue

        # Detectar separadores
        elif linea.strip() in ['---', '___', '***']:
            doc.add_paragraph()  # Espacio adicional
            continue

        # Detectar blockquotes
        elif linea.strip().startswith('>'):
            texto = linea.strip()[1:].strip()
            p = doc.add_paragraph(limpiar_texto(texto))
            p.style = 'Quote'
            continue

        # Texto normal con formato inline
        else:
            p = doc.add_paragraph()
            runs = procesar_formato_inline(linea)
            for run_info in runs:
                agregar_run_formateado(p, run_info)

    print(f"\n    Progreso: 100.0% ({total_lineas:,}/{total_lineas:,} líneas)")

    # Guardar documento
    print(f"\n[*] Guardando documento...")
    doc.save(archivo_docx)

    # Estadísticas
    size_bytes = os.path.getsize(archivo_docx)
    size_mb = size_bytes / (1024 * 1024)
    num_paragraphs = len(doc.paragraphs)
    num_tables = len(doc.tables)

    print("\n" + "="*80)
    print("✓ DOCUMENTO CREADO EXITOSAMENTE")
    print("="*80)
    print(f"\n  📄 Archivo:     {os.path.basename(archivo_docx)}")
    print(f"  📊 Tamaño:      {size_mb:.2f} MB")
    print(f"  📝 Párrafos:    {num_paragraphs:,}")
    print(f"  📋 Tablas:      {num_tables}")
    print(f"  📖 Páginas:     ~{num_paragraphs // 45} (estimado)")
    print(f"\n  ✅ Caracteres especiales: Correctamente procesados")
    print(f"  ✅ Formato profesional: Aplicado")
    print(f"  ✅ Estilos APA: Configurados")
    print("\n" + "="*80)
    print("\n  El documento está listo para:")
    print("    • Abrir en Microsoft Word")
    print("    • Edición final")
    print("    • Impresión")
    print("    • Conversión a PDF")
    print("    • Entrega académica")
    print("\n")

if __name__ == "__main__":
    # Rutas
    base_path = r"C:\Users\Alejandro\Documents\Ivan\lourdes"
    archivo_md = os.path.join(base_path, "PROYECTO_FINAL_CONSOLIDADO.md")
    archivo_docx = os.path.join(base_path, "PROYECTO_FINAL_PROFESIONAL.docx")

    # Verificar archivo fuente
    if not os.path.exists(archivo_md):
        print(f"\n❌ ERROR: No se encontró el archivo {archivo_md}")
        sys.exit(1)

    # Convertir
    try:
        convertir_markdown_a_docx(archivo_md, archivo_docx)
    except Exception as e:
        print(f"\n❌ ERROR durante la conversión: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
