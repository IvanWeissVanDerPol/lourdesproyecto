"""
Comprehensive APA 7 Format Compliance Analyzer
Analyzes a Word document against APA 7 guidelines
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
import re

class APA7Analyzer:
    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.issues = []
        self.warnings = []
        self.passed = []

    def analyze_all(self):
        """Run all APA 7 compliance checks"""
        print("=" * 80)
        print("APA 7 FORMAT COMPLIANCE ANALYSIS")
        print("=" * 80)

        self.check_page_setup()
        self.check_font_and_spacing()
        self.check_title_page()
        self.check_headings()
        self.check_abstract()
        self.check_references()
        self.check_citations()
        self.check_paragraph_formatting()
        self.check_page_numbers()

        self.print_report()

    def check_page_setup(self):
        """Check page margins and size"""
        print("\n[PAGE SETUP] CHECKING PAGE SETUP...")
        section = self.doc.sections[0]

        # Margins should be 1 inch (914400 EMU)
        margin_tolerance = 50000  # Small tolerance
        expected_margin = 914400

        margins = {
            'Top': section.top_margin,
            'Bottom': section.bottom_margin,
            'Left': section.left_margin,
            'Right': section.right_margin
        }

        for margin_name, margin_value in margins.items():
            if abs(margin_value - expected_margin) > margin_tolerance:
                self.issues.append(f"❌ {margin_name} margin is {margin_value/914400:.2f}\" (should be 1\")")
            else:
                self.passed.append(f"✓ {margin_name} margin is correct (1\")")

    def check_font_and_spacing(self):
        """Check font type, size, and line spacing"""
        print("\n[FONT & SPACING] CHECKING FONT AND SPACING...")

        # Acceptable APA 7 fonts
        acceptable_fonts = [
            'Times New Roman',  # 12pt
            'Calibri',          # 11pt
            'Arial',            # 11pt
            'Lucida Sans Unicode',  # 10pt
            'Georgia',          # 11pt
            'Computer Modern'   # 10pt
        ]

        font_sizes = {
            'Times New Roman': 12,
            'Calibri': 11,
            'Arial': 11,
            'Lucida Sans Unicode': 10,
            'Georgia': 11,
            'Computer Modern': 10
        }

        # Check main body paragraphs
        body_paragraphs = [p for p in self.doc.paragraphs if p.text.strip() and
                          p.style.name not in ['Heading 1', 'Heading 2', 'Title']]

        if body_paragraphs:
            sample = body_paragraphs[10] if len(body_paragraphs) > 10 else body_paragraphs[0]

            if sample.runs:
                font_name = sample.runs[0].font.name
                font_size = sample.runs[0].font.size

                if font_name in acceptable_fonts:
                    self.passed.append(f"✓ Font is acceptable: {font_name}")

                    expected_size = Pt(font_sizes[font_name])
                    if font_size and abs(font_size - expected_size) > Pt(0.5):
                        self.issues.append(f"❌ Font size is {font_size.pt}pt, should be {font_sizes[font_name]}pt for {font_name}")
                    else:
                        self.passed.append(f"✓ Font size is correct")
                else:
                    self.warnings.append(f"⚠️  Font '{font_name}' may not be APA 7 compliant")

            # Check line spacing (should be 2.0)
            if sample.paragraph_format.line_spacing:
                if sample.paragraph_format.line_spacing == 2.0:
                    self.passed.append("✓ Line spacing is double (2.0)")
                else:
                    self.issues.append(f"❌ Line spacing is {sample.paragraph_format.line_spacing}, should be 2.0")

    def check_title_page(self):
        """Check title page elements"""
        print("\n[TITLE PAGE] CHECKING TITLE PAGE...")

        # Get first 20 paragraphs (title page area)
        title_area = self.doc.paragraphs[:20]
        text_content = [p.text.strip() for p in title_area if p.text.strip()]

        # Check for required elements
        if len(text_content) >= 3:
            self.passed.append(f"✓ Title page has content: {text_content[0][:60]}...")
        else:
            self.warnings.append("⚠️  Title page may be incomplete")

        # Check for page number on first page
        # (This is complex in python-docx, so we'll note it)
        self.warnings.append("⚠️  Manual check required: Page numbers should start on title page")

    def check_abstract(self):
        """Check abstract section"""
        print("\n[ABSTRACT] CHECKING ABSTRACT...")

        abstract_found = False
        abstract_text = ""

        for i, para in enumerate(self.doc.paragraphs):
            if 'RESUMEN' in para.text.upper() or 'ABSTRACT' in para.text.upper():
                abstract_found = True
                # Get next paragraph as abstract content
                if i + 1 < len(self.doc.paragraphs):
                    abstract_text = self.doc.paragraphs[i + 1].text
                break

        if abstract_found:
            self.passed.append("✓ Abstract/Resumen section found")

            # Check length (150-250 words typically)
            word_count = len(abstract_text.split())
            if 100 <= word_count <= 300:
                self.passed.append(f"✓ Abstract length is appropriate ({word_count} words)")
            elif word_count > 0:
                self.warnings.append(f"⚠️  Abstract may be too {'short' if word_count < 100 else 'long'} ({word_count} words)")
        else:
            self.issues.append("❌ Abstract/Resumen section not found")

        # Check for Keywords
        keywords_found = False
        for para in self.doc.paragraphs:
            if 'PALABRAS CLAVE' in para.text.upper() or 'KEYWORDS' in para.text.upper():
                keywords_found = True
                break

        if keywords_found:
            self.passed.append("✓ Keywords/Palabras Clave section found")
        else:
            self.issues.append("❌ Keywords section not found")

    def check_headings(self):
        """Check heading hierarchy and format"""
        print("\n[HEADINGS] CHECKING HEADINGS...")

        headings = []
        for para in self.doc.paragraphs:
            if 'Heading' in para.style.name:
                headings.append({
                    'level': para.style.name,
                    'text': para.text[:50]
                })

        if headings:
            self.passed.append(f"✓ Found {len(headings)} headings using proper styles")

            # Check for logical hierarchy
            heading_levels = [h['level'] for h in headings[:20]]

            # APA 7 has 5 levels of headings
            unique_levels = set(heading_levels)
            if len(unique_levels) <= 5:
                self.passed.append(f"✓ Using {len(unique_levels)} heading levels (APA allows up to 5)")
            else:
                self.warnings.append(f"⚠️  Using {len(unique_levels)} heading levels, APA recommends max 5")
        else:
            self.warnings.append("⚠️  No formatted headings found (or using manual formatting)")

    def check_references(self):
        """Check references section"""
        print("\n[REFERENCES] CHECKING REFERENCES...")

        ref_section_found = False
        ref_start_idx = -1

        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip().upper() in ['REFERENCIAS', 'REFERENCES', 'REFERENCIAS BIBLIOGRÁFICAS']:
                ref_section_found = True
                ref_start_idx = i
                break

        if ref_section_found:
            self.passed.append("✓ References section found")

            # Check hanging indent in references
            # Get some reference entries
            ref_entries = []
            for i in range(ref_start_idx + 1, min(ref_start_idx + 20, len(self.doc.paragraphs))):
                para = self.doc.paragraphs[i]
                if para.text.strip() and len(para.text) > 30:
                    ref_entries.append(para)

            if ref_entries:
                # Check for hanging indent (first line at 0, subsequent lines indented)
                sample_ref = ref_entries[0]
                if sample_ref.paragraph_format.left_indent and sample_ref.paragraph_format.first_line_indent:
                    if sample_ref.paragraph_format.first_line_indent < 0:
                        self.passed.append("✓ References use hanging indent")
                    else:
                        self.warnings.append("⚠️  References may not have hanging indent")
                else:
                    self.warnings.append("⚠️  Check references for hanging indent (0.5\")")
        else:
            self.issues.append("❌ References section not found")

    def check_citations(self):
        """Check for in-text citations"""
        print("\n[CITATIONS] CHECKING IN-TEXT CITATIONS...")

        # Look for citation patterns
        citation_patterns = [
            r'\([A-Z][a-z]+,\s*\d{4}\)',  # (Author, 2020)
            r'\([A-Z][a-z]+\s*et\s*al\.,\s*\d{4}\)',  # (Author et al., 2020)
            r'\([A-Z][a-z]+\s*&\s*[A-Z][a-z]+,\s*\d{4}\)',  # (Author & Author, 2020)
        ]

        citations_found = 0
        full_text = '\n'.join([p.text for p in self.doc.paragraphs])

        for pattern in citation_patterns:
            matches = re.findall(pattern, full_text)
            citations_found += len(matches)

        if citations_found > 0:
            self.passed.append(f"✓ Found {citations_found} in-text citations")
        else:
            self.warnings.append("⚠️  No parenthetical citations found (or using different format)")

    def check_paragraph_formatting(self):
        """Check paragraph alignment and indentation"""
        print("\n[PARAGRAPHS] CHECKING PARAGRAPH FORMATTING...")

        body_paragraphs = [p for p in self.doc.paragraphs if p.text.strip() and
                          p.style.name not in ['Heading 1', 'Heading 2', 'Title', 'List Bullet', 'List Number']]

        if body_paragraphs:
            sample = body_paragraphs[15] if len(body_paragraphs) > 15 else body_paragraphs[0]

            # Check alignment (should be left-aligned)
            # In python-docx, alignment None means left-aligned
            alignment = sample.paragraph_format.alignment
            if alignment is None or str(alignment) == 'LEFT (0)':
                self.passed.append("✓ Paragraphs are left-aligned")
            else:
                self.warnings.append(f"⚠️  Paragraph alignment may not be standard: {alignment}")

            # Check first line indent (should be 0.5 inches for body text)
            first_line = sample.paragraph_format.first_line_indent
            if first_line:
                expected_indent = Inches(0.5)
                if abs(first_line - expected_indent) < Inches(0.1):
                    self.passed.append("✓ First line indent is correct (0.5\")")
                else:
                    self.warnings.append(f"⚠️  First line indent is {first_line.inches:.2f}\", should be 0.5\"")

    def check_page_numbers(self):
        """Check page numbering"""
        print("\n[PAGE NUMBERS] CHECKING PAGE NUMBERS...")

        # Page numbers are in headers/footers, complex to check in python-docx
        section = self.doc.sections[0]

        if section.header:
            self.warnings.append("⚠️  Manual check: Ensure page numbers are in top right corner")
        else:
            self.warnings.append("⚠️  No header detected - page numbers should be in top right")

    def print_report(self):
        """Print final compliance report"""
        print("\n" + "=" * 80)
        print("COMPLIANCE REPORT SUMMARY")
        print("=" * 80)

        print(f"\n[PASSED] CHECKS PASSED ({len(self.passed)}):")
        for item in self.passed:
            print(f"  {item}")

        print(f"\n[WARNING] WARNINGS ({len(self.warnings)}):")
        if self.warnings:
            for item in self.warnings:
                print(f"  {item}")
        else:
            print("  None")

        print(f"\n[ISSUES] ISSUES FOUND ({len(self.issues)}):")
        if self.issues:
            for item in self.issues:
                print(f"  {item}")
        else:
            print("  None - Great job!")

        # Overall score
        total_checks = len(self.passed) + len(self.warnings) + len(self.issues)
        if total_checks > 0:
            compliance_rate = (len(self.passed) / total_checks) * 100
            print(f"\n[SCORE] COMPLIANCE RATE: {compliance_rate:.1f}%")

            if compliance_rate >= 90:
                print("[EXCELLENT] Document is highly compliant with APA 7!")
            elif compliance_rate >= 75:
                print("[GOOD] Minor improvements needed")
            elif compliance_rate >= 60:
                print("[FAIR] Several issues need attention")
            else:
                print("[NEEDS WORK] Significant formatting issues")

        print("\n" + "=" * 80)

if __name__ == "__main__":
    docx_path = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7.docx"

    analyzer = APA7Analyzer(docx_path)
    analyzer.analyze_all()

    # Detailed structure analysis
    print("\n" + "=" * 80)
    print("DETAILED DOCUMENT STRUCTURE")
    print("=" * 80)

    print(f"\nTotal paragraphs: {len(analyzer.doc.paragraphs)}")
    print(f"Total sections: {len(analyzer.doc.sections)}")

    # Show document outline
    print("\n[OUTLINE] DOCUMENT OUTLINE (First 30 major elements):")
    count = 0
    for i, para in enumerate(analyzer.doc.paragraphs):
        if count >= 30:
            break
        if para.text.strip():
            if 'Heading' in para.style.name:
                print(f"\n{para.style.name.upper()}: {para.text}")
                count += 1
            elif para.style.name == 'Title':
                print(f"\nTITLE: {para.text}")
                count += 1
            elif len(para.text) > 100 and count < 30:
                # Show first paragraph of sections
                print(f"  [{para.style.name}] {para.text[:80]}...")
                count += 1
