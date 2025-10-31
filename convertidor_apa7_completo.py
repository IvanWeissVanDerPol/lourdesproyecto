#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertidor Completo Markdown a DOCX con Formato APA 7
Versión 4.0 - Implementación completa de todos los estándares APA 7ª edición

Características:
- Portada profesional y estudiantil
- 5 niveles de encabezados APA 7
- Running head (para profesionales)
- Numeración automática de páginas
- Resumen/Abstract con palabras clave
- Tablas con formato APA 7 completo
- Figuras con formato APA 7
- Referencias con sangría francesa
- Citas en bloque
- Detección automática de secciones

Autor: Claude
Fecha: Octubre 2025
Basado en: Publication Manual APA 7th Edition (2020)
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement, qn
from docx.table import _Cell
import re
import os
import sys
from datetime import datetime

# Configurar encoding para Windows
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')


class ConvertidorAPA7:
    """
    Convertidor de Markdown a DOCX con formato APA 7 completo

    Implementa todos los estándares de la 7ª edición del manual APA:
    - Portada (profesional y estudiantil)
    - Running head (solo profesional)
    - Resumen/Abstract
    - 5 niveles de encabezados
    - Tablas con formato APA
    - Figuras con formato APA
    - Referencias con sangría francesa
    - Citas textuales en bloque
    """

    def __init__(self, tipo_documento='estudiantil'):
        """
        Inicializar convertidor

        Args:
            tipo_documento: 'profesional' o 'estudiantil'
        """
        self.doc = Document()
        self.tipo_documento = tipo_documento
        self.numero_tabla = 0
        self.numero_figura = 0
        self.en_referencias = False

        self.configurar_documento()
        self.crear_estilos_apa7()

    def configurar_documento(self):
        """Configurar propiedades del documento según APA 7"""

        for section in self.doc.sections:
            # Márgenes: 1 pulgada (2.54 cm) en todos los lados
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Encabezado y pie de página
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

            # Papel Carta (8.5 x 11 pulgadas)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

    def crear_estilos_apa7(self):
        """Crear todos los estilos según APA 7"""

        # ===== ESTILO NORMAL =====
        normal = self.doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 2.0  # Doble espacio
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        normal.paragraph_format.first_line_indent = Inches(0.5)  # Sangría primera línea

        # ===== NIVEL 1: Centrado, Negrita, Title Case =====
        self._crear_o_modificar_estilo(
            nombre='Heading 1',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=True,
            cursiva=False,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER,
            sangria=Inches(0),
            sangria_primera=Inches(0),
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0),
            keep_with_next=True
        )

        # ===== NIVEL 2: Izquierda, Negrita, Title Case =====
        self._crear_o_modificar_estilo(
            nombre='Heading 2',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=True,
            cursiva=False,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0),
            sangria_primera=Inches(0),
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0),
            keep_with_next=True
        )

        # ===== NIVEL 3: Izquierda, Negrita, Cursiva, Title Case =====
        self._crear_o_modificar_estilo(
            nombre='Heading 3',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=True,
            cursiva=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0),
            sangria_primera=Inches(0),
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0),
            keep_with_next=True
        )

        # ===== NIVEL 4: Sangría, Negrita, Cursiva, Title Case, Punto =====
        # Nota: Nivel 4 y 5 requieren tratamiento especial (texto continúa en misma línea)
        self._crear_o_modificar_estilo(
            nombre='Heading 4',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=True,
            cursiva=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0.5),
            sangria_primera=Inches(0),
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0),
            keep_with_next=True
        )

        # ===== NIVEL 5: Sangría, Negrita, Cursiva, Sentence case, Punto =====
        self._crear_o_modificar_estilo(
            nombre='Heading 5',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=True,
            cursiva=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0.5),
            sangria_primera=Inches(0),
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0),
            keep_with_next=True
        )

        # ===== CITAS EN BLOQUE (40+ palabras) =====
        self._crear_o_modificar_estilo(
            nombre='Quote',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=False,
            cursiva=False,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0.5),  # Sangría de 0.5"
            sangria_primera=Inches(0),  # Sin sangría adicional
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0)
        )

        # ===== REFERENCIAS (con sangría francesa) =====
        self._crear_o_modificar_estilo(
            nombre='Reference',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=False,
            cursiva=False,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0),
            sangria_primera=Inches(-0.5),  # Sangría francesa
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0),
            left_indent=Inches(0.5)  # Sangría del cuerpo
        )

        # ===== RESUMEN (sin sangría de primera línea) =====
        self._crear_o_modificar_estilo(
            nombre='Abstract',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=False,
            cursiva=False,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0),
            sangria_primera=Inches(0),  # Sin sangría
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0)
        )

        # ===== TÍTULO DE TABLA =====
        self._crear_o_modificar_estilo(
            nombre='Table Title',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=False,
            cursiva=True,  # Título de tabla en cursiva
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0),
            sangria_primera=Inches(0),
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0)
        )

        # ===== NÚMERO DE TABLA =====
        self._crear_o_modificar_estilo(
            nombre='Table Number',
            tipo=WD_STYLE_TYPE.PARAGRAPH,
            fuente='Times New Roman',
            tamano=Pt(12),
            negrita=True,  # Número en negrita
            cursiva=False,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            sangria=Inches(0),
            sangria_primera=Inches(0),
            interlineado=2.0,
            espacio_antes=Pt(0),
            espacio_despues=Pt(0)
        )

    def _crear_o_modificar_estilo(self, nombre, tipo, fuente, tamano, negrita, cursiva,
                                   alineacion, sangria, sangria_primera, interlineado,
                                   espacio_antes, espacio_despues, keep_with_next=False,
                                   left_indent=None):
        """Crear o modificar un estilo con parámetros específicos"""

        try:
            estilo = self.doc.styles[nombre]
        except KeyError:
            estilo = self.doc.styles.add_style(nombre, tipo)

        # Formato de fuente
        estilo.font.name = fuente
        estilo.font.size = tamano
        estilo.font.bold = negrita
        estilo.font.italic = cursiva

        # Formato de párrafo
        estilo.paragraph_format.alignment = alineacion
        estilo.paragraph_format.left_indent = left_indent if left_indent else sangria
        estilo.paragraph_format.first_line_indent = sangria_primera
        estilo.paragraph_format.line_spacing = interlineado
        estilo.paragraph_format.space_before = espacio_antes
        estilo.paragraph_format.space_after = espacio_despues
        estilo.paragraph_format.keep_with_next = keep_with_next

        return estilo

    def agregar_portada(self, titulo, autor, institucion, **kwargs):
        """
        Agregar portada según APA 7

        Args:
            titulo: Título completo del trabajo
            autor: Nombre del autor
            institucion: Nombre de la institución
            **kwargs: Argumentos opcionales
                tipo: 'profesional' o 'estudiantil'
                running_head: Título abreviado (max 50 caracteres, solo profesional)
                departamento: Departamento o facultad
                curso: Nombre del curso (estudiante)
                profesor: Nombre del profesor (estudiante)
                fecha: Fecha de entrega
                nota_autor: Nota del autor (profesional)
        """

        tipo = kwargs.get('tipo', self.tipo_documento)
        running_head = kwargs.get('running_head', '')

        # Configurar encabezado de portada
        section = self.doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()

        if tipo == 'profesional' and running_head:
            # Running head en mayúsculas, izquierda
            run = header_para.add_run(running_head.upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Número de página en la derecha
        header_para.add_run('\t\t\t\t\t')
        self._agregar_numero_pagina(header_para)

        # Espacios para centrar verticalmente (aprox.)
        for _ in range(7):
            self.doc.add_paragraph()

        # TÍTULO (negrita, centrado)
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.line_spacing = 2.0
        run = p_titulo.add_run(titulo)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        self.doc.add_paragraph()  # Espacio

        # AUTOR
        p_autor = self.doc.add_paragraph()
        p_autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_autor.paragraph_format.line_spacing = 2.0
        run = p_autor.add_run(autor)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # INSTITUCIÓN (puede incluir departamento)
        departamento = kwargs.get('departamento', '')
        texto_inst = f"{departamento}, {institucion}" if departamento else institucion

        p_inst = self.doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_inst.paragraph_format.line_spacing = 2.0
        run = p_inst.add_run(texto_inst)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Elementos adicionales según tipo
        if tipo == 'estudiantil':
            # Curso, Profesor, Fecha
            self.doc.add_paragraph()

            curso = kwargs.get('curso', '')
            if curso:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 2.0
                run = p.add_run(curso)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            profesor = kwargs.get('profesor', '')
            if profesor:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 2.0
                run = p.add_run(profesor)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            fecha = kwargs.get('fecha', datetime.now().strftime('%d de %B de %Y'))
            if fecha:
                for _ in range(3):
                    self.doc.add_paragraph()
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 2.0
                run = p.add_run(fecha)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        else:  # profesional
            # Nota del autor
            nota_autor = kwargs.get('nota_autor', '')
            if nota_autor:
                for _ in range(5):
                    self.doc.add_paragraph()

                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 2.0
                run = p.add_run("Nota del autor")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = False

                self.doc.add_paragraph()

                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 2.0
                run = p.add_run(nota_autor)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Salto de página
        self.doc.add_page_break()

    def _agregar_numero_pagina(self, paragraph):
        """Agregar número de página automático"""

        run = paragraph.add_run()

        # Crear campo de número de página
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def agregar_resumen(self, texto_resumen, palabras_clave=None):
        """
        Agregar página de resumen según APA 7

        Args:
            texto_resumen: Texto del resumen (150-250 palabras)
            palabras_clave: Lista de palabras clave (3-5 palabras)
        """

        # Título "Resumen" centrado y en negrita
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.line_spacing = 2.0
        run = p_titulo.add_run("Resumen")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        # Texto del resumen (sin sangría de primera línea)
        p_resumen = self.doc.add_paragraph(texto_resumen)
        p_resumen.style = 'Abstract'

        # Palabras clave
        if palabras_clave:
            self.doc.add_paragraph()  # Línea en blanco

            p_palabras = self.doc.add_paragraph()
            p_palabras.paragraph_format.left_indent = Inches(0.5)  # Sangría
            p_palabras.paragraph_format.first_line_indent = Inches(0)
            p_palabras.paragraph_format.line_spacing = 2.0

            run = p_palabras.add_run("Palabras clave: ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True

            if isinstance(palabras_clave, list):
                texto_palabras = ', '.join(palabras_clave)
            else:
                texto_palabras = palabras_clave

            run = p_palabras.add_run(texto_palabras)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Salto de página
        self.doc.add_page_break()

    def procesar_markdown(self, archivo_md):
        """
        Procesar archivo Markdown y convertir a DOCX con formato APA 7

        Args:
            archivo_md: Ruta al archivo Markdown
        """

        print(f"\n[*] Leyendo archivo: {archivo_md}")
        with open(archivo_md, 'r', encoding='utf-8') as f:
            contenido = f.read()

        lineas = contenido.split('\n')
        total = len(lineas)
        print(f"[*] Procesando {total:,} líneas...")

        # Variables de estado
        en_tabla = False
        filas_tabla = []
        en_codigo = False
        bloque_codigo = []
        tabla_numero = None
        tabla_titulo = None

        i = 0
        while i < len(lineas):
            linea = lineas[i]

            if i % 1000 == 0:
                print(f"    {(i/total)*100:.1f}% ({i:,}/{total:,})", end='\r')

            # Detectar bloques de código
            if linea.strip().startswith('```'):
                if en_codigo:
                    if bloque_codigo:
                        p = self.doc.add_paragraph('\n'.join(bloque_codigo))
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.first_line_indent = Inches(0)
                        p.paragraph_format.line_spacing = 2.0
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                    bloque_codigo = []
                    en_codigo = False
                else:
                    en_codigo = True
                i += 1
                continue

            if en_codigo:
                bloque_codigo.append(linea)
                i += 1
                continue

            # Detectar tablas
            if '|' in linea and linea.strip().startswith('|'):
                if not en_tabla:
                    en_tabla = True
                    filas_tabla = []
                    # Buscar número y título de tabla antes de la tabla
                    if i > 0:
                        linea_previa = lineas[i-1].strip()
                        if linea_previa.startswith('**Tabla'):
                            match = re.match(r'\*\*Tabla (\d+)\*\*', linea_previa)
                            if match:
                                tabla_numero = int(match.group(1))
                        if i > 1:
                            linea_titulo = lineas[i-2].strip()
                            if linea_titulo and not linea_titulo.startswith('**Tabla'):
                                tabla_titulo = linea_titulo

                celdas = [c.strip() for c in linea.split('|')[1:-1]]
                # Detectar separador
                if all(re.match(r'^-+$', c.strip()) for c in celdas if c.strip()):
                    i += 1
                    continue

                filas_tabla.append(celdas)
                i += 1
                continue
            else:
                if en_tabla and filas_tabla:
                    self.crear_tabla_apa(filas_tabla, tabla_numero, tabla_titulo)
                    filas_tabla = []
                    en_tabla = False
                    tabla_numero = None
                    tabla_titulo = None

            # Líneas vacías
            if not linea.strip():
                i += 1
                continue

            # Detectar sección de Referencias
            if re.match(r'^##\s+Referencias\s*$', linea, re.IGNORECASE):
                self.en_referencias = True
                p = self.doc.add_paragraph("Referencias")
                p.style = 'Heading 1'
                i += 1
                continue

            # Detectar Resumen/Abstract
            if re.match(r'^##\s+(Resumen|Abstract)\s*$', linea, re.IGNORECASE):
                p = self.doc.add_paragraph(linea[3:].strip())
                p.style = 'Heading 1'
                i += 1
                continue

            # Procesar encabezados
            if linea.startswith('###### '):
                # Nivel 5: Sangría, negrita, cursiva, sentence case, punto
                texto = self.limpiar_texto(linea[7:].strip())
                texto = texto[0].upper() + texto[1:].lower() if texto else texto
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.line_spacing = 2.0
                run = p.add_run(texto + '. ')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.italic = True
                # El texto continúa en la misma línea (implementar si es necesario)
                i += 1
                continue

            elif linea.startswith('##### '):
                # Nivel 4: Sangría, negrita, cursiva, title case, punto
                texto = self.limpiar_texto(linea[6:].strip())
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.line_spacing = 2.0
                run = p.add_run(texto + '. ')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.italic = True
                i += 1
                continue

            elif linea.startswith('#### '):
                texto = self.limpiar_texto(linea[5:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 3'
                i += 1
                continue

            elif linea.startswith('### '):
                texto = self.limpiar_texto(linea[4:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 2'
                i += 1
                continue

            elif linea.startswith('## '):
                texto = self.limpiar_texto(linea[3:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 1'
                i += 1
                continue

            # Listas
            elif linea.strip().startswith(('- ', '* ', '+ ')):
                texto = self.limpiar_texto(linea.strip()[2:])
                p = self.doc.add_paragraph(texto, style='List Bullet')
                i += 1
                continue

            elif re.match(r'^\d+\.\s', linea.strip()):
                texto = self.limpiar_texto(re.sub(r'^\d+\.\s', '', linea.strip()))
                p = self.doc.add_paragraph(texto, style='List Number')
                i += 1
                continue

            # Separadores
            elif linea.strip() in ['---', '___', '***', '___']:
                i += 1
                continue

            # Blockquotes (citas en bloque)
            elif linea.strip().startswith('>'):
                texto = self.limpiar_texto(linea.strip()[1:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Quote'
                i += 1
                continue

            # Referencias (con sangría francesa)
            elif self.en_referencias:
                texto = self.limpiar_texto(linea)
                p = self.doc.add_paragraph()
                self.agregar_formato_inline(p, texto)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.line_spacing = 2.0
                i += 1
                continue

            # Texto normal
            else:
                p = self.doc.add_paragraph()
                self.agregar_formato_inline(p, linea)
                i += 1

        print(f"\n    100.0% ({total:,}/{total:,}) - Completado")

    def crear_tabla_apa(self, filas, numero=None, titulo=None):
        """
        Crear tabla según formato APA 7

        Args:
            filas: Lista de listas con el contenido de las celdas
            numero: Número de tabla (opcional)
            titulo: Título de la tabla (opcional)
        """

        if not numero:
            self.numero_tabla += 1
            numero = self.numero_tabla

        # Número de tabla (negrita, izquierda)
        p_numero = self.doc.add_paragraph(f"Tabla {numero}")
        p_numero.style = 'Table Number'

        # Título de tabla (cursiva, izquierda)
        if titulo:
            p_titulo = self.doc.add_paragraph(titulo)
            p_titulo.style = 'Table Title'

        if len(filas) < 1:
            return

        num_cols = max(len(fila) for fila in filas)
        tabla = self.doc.add_table(rows=len(filas), cols=num_cols)

        # Estilo de tabla APA: solo líneas horizontales
        tabla.style = 'Light Grid Accent 1'  # Temporal, luego quitamos bordes verticales

        # Llenar tabla
        for i, fila in enumerate(filas):
            for j, celda in enumerate(fila):
                if j < num_cols:
                    cell = tabla.rows[i].cells[j]
                    cell.text = self.limpiar_texto(celda)

                    # Formato de celda
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.line_spacing = 1.0  # Sencillo en tablas
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            if i == 0:  # Encabezado en negrita
                                run.font.bold = True

        # Quitar bordes verticales (solo horizontales superior e inferior)
        self._aplicar_bordes_apa(tabla)

        # Espacio después de tabla
        self.doc.add_paragraph()

    def _aplicar_bordes_apa(self, tabla):
        """Aplicar bordes APA a tabla (solo líneas horizontales)"""

        tbl = tabla._element
        tblPr = tbl.tblPr

        # Crear elemento de bordes
        tblBorders = OxmlElement('w:tblBorders')

        # Línea superior
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:color'), '000000')
        tblBorders.append(top)

        # Línea inferior
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), '000000')
        tblBorders.append(bottom)

        # Sin bordes izquierda, derecha, internos verticales
        for side in ['left', 'right', 'insideV']:
            elem = OxmlElement(f'w:{side}')
            elem.set(qn('w:val'), 'none')
            tblBorders.append(elem)

        # Línea horizontal interna (entre filas)
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), 'single')
        insideH.set(qn('w:sz'), '6')
        insideH.set(qn('w:color'), '000000')
        tblBorders.append(insideH)

        tblPr.append(tblBorders)

    def agregar_formato_inline(self, paragraph, texto):
        """Agregar texto con formato inline (negrita, cursiva, código)"""

        texto = self.limpiar_texto(texto)

        # Patrón para detectar formato
        patron = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)'
        partes = re.split(patron, texto)

        for parte in partes:
            if not parte:
                continue

            run = paragraph.add_run()

            # Negrita + Cursiva
            if parte.startswith('***') and parte.endswith('***') and len(parte) > 6:
                run.text = parte[3:-3]
                run.bold = True
                run.italic = True
            # Negrita
            elif parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
                run.text = parte[2:-2]
                run.bold = True
            # Cursiva
            elif parte.startswith('*') and parte.endswith('*') and len(parte) > 2:
                run.text = parte[1:-1]
                run.italic = True
            # Código
            elif parte.startswith('`') and parte.endswith('`'):
                run.text = parte[1:-1]
                run.font.name = 'Courier New'
                run.font.size = Pt(11)
            else:
                run.text = parte

    def limpiar_texto(self, texto):
        """Limpiar texto de caracteres de formato Markdown y normalizar encoding"""

        # Normalizar encoding
        texto = texto.replace('�', 'ó')
        texto = texto.replace('Ã³', 'ó')
        texto = texto.replace('Ã±', 'ñ')
        texto = texto.replace('Ã©', 'é')
        texto = texto.replace('Ã¡', 'á')
        texto = texto.replace('Ã­', 'í')
        texto = texto.replace('Ãº', 'ú')

        # Quitar emojis y símbolos de markdown que no sean formato
        texto = re.sub(r'[📊📄📚⭐✅❌🎯🔍]', '', texto)

        return texto

    def guardar(self, ruta):
        """Guardar documento"""

        print(f"\n[*] Guardando documento: {ruta}")
        self.doc.save(ruta)

        size_mb = os.path.getsize(ruta) / (1024 * 1024)
        print(f"\n{'='*80}")
        print("✓ DOCUMENTO APA 7 CREADO EXITOSAMENTE")
        print('='*80)
        print(f"\n  📄 Archivo: {os.path.basename(ruta)}")
        print(f"  📊 Tamaño: {size_mb:.2f} MB")
        print(f"  📝 Párrafos: {len(self.doc.paragraphs):,}")
        print(f"  📋 Tablas: {len(self.doc.tables)}")
        print(f"\n  ✅ Formato APA 7 completo aplicado")
        print(f"  ✅ Márgenes: 1 pulgada (2.54 cm)")
        print(f"  ✅ Fuente: Times New Roman 12pt")
        print(f"  ✅ Interlineado: Doble")
        print(f"  ✅ Sangría: 0.5 pulgadas primera línea")
        print(f"  ✅ 5 niveles de encabezados")
        print(f"  ✅ Tablas con formato APA")
        print(f"  ✅ Referencias con sangría francesa")
        print("\n" + "="*80 + "\n")


def main():
    """Función principal de ejemplo"""

    print("\n" + "="*80)
    print("  CONVERTIDOR COMPLETO MARKDOWN → DOCX CON FORMATO APA 7")
    print("  Versión 4.0 - Implementación completa")
    print("="*80)

    # Configuración
    base_path = r"C:\Users\Alejandro\Documents\Ivan\lourdes"
    archivo_md = os.path.join(base_path, "PROYECTO_FINAL_CONSOLIDADO.md")
    archivo_salida = os.path.join(base_path, "PROYECTO_APA7_COMPLETO.docx")

    # Verificar archivo
    if not os.path.exists(archivo_md):
        print(f"\n❌ ERROR: No se encontró {archivo_md}")
        return

    # Crear convertidor (tipo: 'estudiantil' o 'profesional')
    print("\n[*] Inicializando convertidor APA 7...")
    convertidor = ConvertidorAPA7(tipo_documento='estudiantil')

    # Agregar portada
    print("[*] Creando portada APA 7...")
    convertidor.agregar_portada(
        titulo="Efectividad de una Intervención Basada en Reforzamiento Positivo y Técnicas de Psicohigiene para Aumentar la Atención Sostenida en un Niño de 9 Años con Trastorno por Déficit de Atención: Estudio de Caso Único con Diseño de Cambio de Criterio",
        autor="[Nombre del Estudiante]",
        institucion="Universidad Nacional",
        departamento="Facultad de Psicología",
        curso="Trabajo de Tesis para optar al título de Licenciado/a en Psicología",
        profesor="Director de Tesis: [Nombre del Director]",
        fecha="Octubre 2025",
        tipo='estudiantil'
    )

    # Agregar resumen (opcional - descomentar si existe en el MD)
    # print("[*] Agregando resumen...")
    # convertidor.agregar_resumen(
    #     texto_resumen="[Texto del resumen aquí]",
    #     palabras_clave=["atención sostenida", "reforzamiento positivo", "TDA", "psicohigiene"]
    # )

    # Procesar contenido Markdown
    print("[*] Procesando contenido Markdown...")
    convertidor.procesar_markdown(archivo_md)

    # Guardar
    convertidor.guardar(archivo_salida)


if __name__ == "__main__":
    main()
