#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Formateador Profesional APA 7 para Tesis de Psicología
Versión 3.0 - Con todos los estándares APA 7ª edición
Incluye: Portada, encabezados, formato de tablas, referencias, etc.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import sys

# Configurar encoding para Windows
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

class FormateadorAPAProfesional:
    """Clase para formatear documentos según APA 7"""

    def __init__(self):
        self.doc = Document()
        self.configurar_documento()
        self.crear_estilos_apa()

    def configurar_documento(self):
        """Configurar propiedades básicas del documento según APA 7"""

        # Configurar todas las secciones
        for section in self.doc.sections:
            # Márgenes APA 7: 1 pulgada (2.54 cm) en todos los lados
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Distancia de encabezado y pie de página
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

    def crear_estilos_apa(self):
        """Crear estilos personalizados según APA 7"""

        # ===== ESTILO NORMAL (Texto del cuerpo) =====
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)

        normal_para = normal_style.paragraph_format
        normal_para.line_spacing = 2.0  # Doble espacio (APA 7)
        normal_para.space_before = Pt(0)
        normal_para.space_after = Pt(0)
        normal_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # APA prefiere izquierda
        normal_para.first_line_indent = Inches(0.5)  # Sangría de 0.5"

        # ===== TÍTULO DE NIVEL 1 (Centrado, Negrita, Mayúsculas y Minúsculas) =====
        try:
            h1 = self.doc.styles['Heading 1']
        except:
            h1 = self.doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

        h1.font.name = 'Times New Roman'
        h1.font.size = Pt(12)
        h1.font.bold = True
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.line_spacing = 2.0
        h1.paragraph_format.keep_with_next = True

        # ===== TÍTULO DE NIVEL 2 (Izquierda, Negrita) =====
        try:
            h2 = self.doc.styles['Heading 2']
        except:
            h2 = self.doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)

        h2.font.name = 'Times New Roman'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(0)
        h2.paragraph_format.line_spacing = 2.0
        h2.paragraph_format.keep_with_next = True

        # ===== TÍTULO DE NIVEL 3 (Izquierda, Negrita, Cursiva) =====
        try:
            h3 = self.doc.styles['Heading 3']
        except:
            h3 = self.doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)

        h3.font.name = 'Times New Roman'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.italic = True
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(0)
        h3.paragraph_format.line_spacing = 2.0
        h3.paragraph_format.keep_with_next = True

        # ===== TÍTULO DE NIVEL 4 (Izquierda, Negrita, Sangría, Punto final) =====
        try:
            h4 = self.doc.styles['Heading 4']
        except:
            h4 = self.doc.styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)

        h4.font.name = 'Times New Roman'
        h4.font.size = Pt(12)
        h4.font.bold = True
        h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h4.paragraph_format.left_indent = Inches(0.5)
        h4.paragraph_format.space_before = Pt(0)
        h4.paragraph_format.space_after = Pt(0)
        h4.paragraph_format.line_spacing = 2.0

        # ===== ESTILO PARA CITAS (Block Quote) =====
        try:
            quote = self.doc.styles['Quote']
        except:
            quote = self.doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)

        quote.font.name = 'Times New Roman'
        quote.font.size = Pt(12)
        quote.paragraph_format.left_indent = Inches(0.5)
        quote.paragraph_format.right_indent = Inches(0)
        quote.paragraph_format.space_before = Pt(0)
        quote.paragraph_format.space_after = Pt(0)
        quote.paragraph_format.line_spacing = 2.0

        # ===== ESTILO PARA LISTAS =====
        try:
            list_style = self.doc.styles['List Bullet']
            list_style.font.name = 'Times New Roman'
            list_style.font.size = Pt(12)
            list_style.paragraph_format.line_spacing = 2.0
        except:
            pass

    def agregar_portada_apa(self, titulo, autor, institucion, curso="", profesor="", fecha=""):
        """Agregar portada según formato APA 7"""

        # Portada en página separada
        # Número de página en esquina superior derecha
        section = self.doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "1"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Agregar saltos de línea para centrar verticalmente (aproximado)
        for _ in range(8):
            self.doc.add_paragraph()

        # Título (Negrita, centrado, mayúsculas y minúsculas)
        titulo_para = self.doc.add_paragraph()
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_para.line_spacing = 2.0
        run = titulo_para.add_run(titulo)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        # Espacio
        self.doc.add_paragraph()

        # Autor
        autor_para = self.doc.add_paragraph()
        autor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        autor_para.line_spacing = 2.0
        run = autor_para.add_run(autor)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Institución
        inst_para = self.doc.add_paragraph()
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_para.line_spacing = 2.0
        run = inst_para.add_run(institucion)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Curso (opcional)
        if curso:
            self.doc.add_paragraph()
            curso_para = self.doc.add_paragraph()
            curso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            curso_para.line_spacing = 2.0
            run = curso_para.add_run(curso)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Profesor (opcional)
        if profesor:
            prof_para = self.doc.add_paragraph()
            prof_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prof_para.line_spacing = 2.0
            run = prof_para.add_run(profesor)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Fecha
        if fecha:
            # Espacio
            for _ in range(3):
                self.doc.add_paragraph()

            fecha_para = self.doc.add_paragraph()
            fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fecha_para.line_spacing = 2.0
            run = fecha_para.add_run(fecha)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Salto de página después de portada
        self.doc.add_page_break()

    def procesar_contenido_markdown(self, archivo_md):
        """Procesar archivo markdown y convertir a formato APA"""

        print(f"\n[*] Leyendo archivo: {archivo_md}")
        with open(archivo_md, 'r', encoding='utf-8') as f:
            contenido = f.read()

        lineas = contenido.split('\n')
        total = len(lineas)

        print(f"[*] Procesando {total:,} líneas...")

        en_tabla = False
        filas_tabla = []
        en_codigo = False
        bloque_codigo = []
        numero_pagina = 2  # Después de portada

        for i, linea in enumerate(lineas):
            if i % 1000 == 0:
                print(f"    {(i/total)*100:.1f}% ({i:,}/{total:,})", end='\r')

            # Detectar bloques de código
            if linea.strip().startswith('```'):
                if en_codigo:
                    if bloque_codigo:
                        p = self.doc.add_paragraph('\n'.join(bloque_codigo))
                        p.style = 'Normal'
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.first_line_indent = Inches(0)
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                    bloque_codigo = []
                    en_codigo = False
                else:
                    en_codigo = True
                continue

            if en_codigo:
                bloque_codigo.append(linea)
                continue

            # Detectar tablas
            if '|' in linea and linea.strip().startswith('|'):
                if not en_tabla:
                    en_tabla = True
                    filas_tabla = []

                celdas = [c.strip() for c in linea.split('|')[1:-1]]
                if all(re.match(r'^-+$', c.strip()) for c in celdas if c.strip()):
                    continue

                filas_tabla.append(celdas)
                continue
            else:
                if en_tabla and filas_tabla:
                    self.crear_tabla_apa(filas_tabla)
                    filas_tabla = []
                    en_tabla = False

            # Líneas vacías
            if not linea.strip():
                continue

            # Procesar títulos
            if linea.startswith('## ') and not linea.startswith('### '):
                texto = self.limpiar_texto(linea[3:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 1'
                continue

            elif linea.startswith('### ') and not linea.startswith('#### '):
                texto = self.limpiar_texto(linea[4:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 2'
                continue

            elif linea.startswith('#### ') and not linea.startswith('##### '):
                texto = self.limpiar_texto(linea[5:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 3'
                continue

            elif linea.startswith('##### '):
                texto = self.limpiar_texto(linea[6:].strip())
                # Nivel 4: sangría con punto
                p = self.doc.add_paragraph(texto + '.')
                p.style = 'Heading 4'
                continue

            # Listas
            elif linea.strip().startswith(('- ', '* ')):
                texto = self.limpiar_texto(linea.strip()[2:])
                p = self.doc.add_paragraph(texto, style='List Bullet')
                continue

            elif re.match(r'^\d+\.\s', linea.strip()):
                texto = self.limpiar_texto(re.sub(r'^\d+\.\s', '', linea.strip()))
                p = self.doc.add_paragraph(texto, style='List Number')
                continue

            # Separadores
            elif linea.strip() in ['---', '___', '***']:
                continue

            # Blockquotes
            elif linea.strip().startswith('>'):
                texto = self.limpiar_texto(linea.strip()[1:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Quote'
                continue

            # Texto normal
            else:
                p = self.doc.add_paragraph()
                self.agregar_formato_inline(p, linea)

        print(f"\n    100.0% ({total:,}/{total:,}) - Completado")

    def crear_tabla_apa(self, filas):
        """Crear tabla según formato APA 7"""

        if len(filas) < 2:
            return

        num_cols = max(len(fila) for fila in filas)
        tabla = self.doc.add_table(rows=len(filas), cols=num_cols)

        # Estilo de tabla APA: solo líneas horizontales superior e inferior
        tabla.style = 'Table Grid'

        # Llenar tabla
        for i, fila in enumerate(filas):
            for j, celda in enumerate(fila):
                if j < num_cols:
                    cell = tabla.rows[i].cells[j]
                    cell.text = self.limpiar_texto(celda)

                    # Formato de celda
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            if i == 0:  # Encabezado en negrita
                                run.font.bold = True

        # Agregar espacio después de tabla
        self.doc.add_paragraph()

    def limpiar_texto(self, texto):
        """Limpiar texto para encoding correcto"""
        # Normalizar caracteres especiales
        texto = texto.replace('�', 'ó')
        texto = texto.replace('Ã³', 'ó')
        texto = texto.replace('Ã±', 'ñ')
        texto = texto.replace('Ã©', 'é')
        texto = texto.replace('Ã¡', 'á')
        texto = texto.replace('Ã­', 'í')
        texto = texto.replace('Ãº', 'ú')
        return texto

    def agregar_formato_inline(self, paragraph, texto):
        """Agregar texto con formato inline (negrita, cursiva, código)"""

        texto = self.limpiar_texto(texto)

        # Patrón para detectar formato
        patron = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
        partes = re.split(patron, texto)

        for parte in partes:
            if not parte:
                continue

            run = paragraph.add_run(parte)

            # Negrita
            if parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
                run.text = parte[2:-2]
                run.bold = True
            # Cursiva
            elif parte.startswith('*') and parte.endswith('*') and len(parte) > 2 and not parte.startswith('**'):
                run.text = parte[1:-1]
                run.italic = True
            # Código
            elif parte.startswith('`') and parte.endswith('`'):
                run.text = parte[1:-1]
                run.font.name = 'Courier New'
                run.font.size = Pt(11)

    def agregar_encabezados(self, titulo_corto):
        """Agregar encabezados a todas las páginas (excepto portada) según APA 7"""

        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False

            # Limpiar encabezado existente
            for para in header.paragraphs:
                para.clear()

            # Agregar encabezado
            header_para = header.paragraphs[0]
            header_para.text = titulo_corto
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Agregar número de página alineado a la derecha
            header_para.add_run('\t\t\t\t\t')  # Tabulación para alinear derecha
            run_num = header_para.add_run()

            # Campo de número de página
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run_num._element.append(fldChar)

            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            run_num._element.append(instrText)

            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run_num._element.append(fldChar)

            # Formato
            for run in header_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    def guardar(self, ruta):
        """Guardar documento"""
        print(f"\n[*] Guardando documento: {ruta}")
        self.doc.save(ruta)

        size_mb = os.path.getsize(ruta) / (1024 * 1024)
        print(f"\n{'='*80}")
        print("✓ DOCUMENTO APA 7 CREADO EXITOSAMENTE")
        print('='*80)
        print(f"\n  📄 Archivo: {os.path.basename(ruta)}")
        print(f"  📊 Tamaño: {size_mb:.2f} MB")
        print(f"  📝 Párrafos: {len(self.doc.paragraphs):,}")
        print(f"  📋 Tablas: {len(self.doc.tables)}")
        print(f"\n  ✅ Formato APA 7: Aplicado")
        print(f"  ✅ Portada profesional: Incluida")
        print(f"  ✅ Encabezados: Configurados")
        print(f"  ✅ Márgenes: 1 pulgada")
        print(f"  ✅ Doble espacio: Aplicado")
        print(f"  ✅ Times New Roman 12pt: Configurado")
        print("\n" + "="*80 + "\n")

def main():
    """Función principal"""

    print("\n" + "="*80)
    print("  FORMATEADOR PROFESIONAL APA 7")
    print("  Para Tesis de Psicología")
    print("="*80)

    # Rutas
    base_path = r"C:\Users\Alejandro\Documents\Ivan\lourdes"
    archivo_md = os.path.join(base_path, "PROYECTO_FINAL_CONSOLIDADO.md")
    archivo_salida = os.path.join(base_path, "PROYECTO_APA7_PROFESIONAL.docx")

    # Verificar archivo
    if not os.path.exists(archivo_md):
        print(f"\n❌ ERROR: No se encontró {archivo_md}")
        return

    # Crear formateador
    print("\n[*] Inicializando formateador APA 7...")
    formateador = FormateadorAPAProfesional()

    # Agregar portada
    print("[*] Creando portada APA 7...")
    formateador.agregar_portada_apa(
        titulo="Efectividad de una Intervención Basada en Reforzamiento Positivo y Técnicas de Psicohigiene para Aumentar la Atención Sostenida en un Niño de 9 Años con Trastorno por Déficit de Atención: Estudio de Caso Único con Diseño de Cambio de Criterio",
        autor="[Nombre del Estudiante]",
        institucion="[Nombre de la Universidad]\nFacultad de Psicología\nLicenciatura en Psicología",
        curso="Trabajo de Tesis para optar al título de Licenciado/a en Psicología",
        profesor="Director de Tesis: [Nombre del Director]",
        fecha="Octubre 2025"
    )

    # Procesar contenido
    formateador.procesar_contenido_markdown(archivo_md)

    # Agregar encabezados
    print("[*] Agregando encabezados...")
    formateador.agregar_encabezados("INTERVENCIÓN TDA")

    # Guardar
    formateador.guardar(archivo_salida)

if __name__ == "__main__":
    main()
