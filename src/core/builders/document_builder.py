#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Builder - Constructor principal de documentos DOCX

Este módulo orquesta la construcción del documento DOCX a partir
de los elementos parseados del Markdown.
"""

from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

from ..parsers.markdown_parser import MarkdownElement, ElementType
from ..parsers.inline_formatter import InlineFormatter, InlineFormat
from ..styles.style_factory import StyleFactory
from ..utils.exceptions import BuildError
from ...config.apa7_config import APAConfig

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """
    Constructor principal de documentos DOCX con formato APA 7

    Convierte una lista de MarkdownElement en un documento DOCX
    completo con todos los estilos APA 7 aplicados.
    """

    def __init__(self, config: APAConfig):
        """
        Inicializar builder

        Args:
            config: Configuración APA 7
        """
        self.config = config
        self.document = Document()
        self.style_factory = StyleFactory(self.document)
        self.inline_formatter = InlineFormatter()
        self.in_references = False
        self.table_counter = 0
        self.figure_counter = 0

        # Configurar documento
        self._configure_document()
        # Aplicar estilos APA
        self.style_factory.apply_all_apa_styles()

    def _configure_document(self) -> None:
        """Configurar propiedades del documento"""
        for section in self.document.sections:
            # Márgenes
            margin_value = self.config.margins['top']
            unit = self.config.margins.get('units', 'inches')

            if unit == 'inches':
                section.top_margin = Inches(margin_value)
                section.bottom_margin = Inches(self.config.margins['bottom'])
                section.left_margin = Inches(self.config.margins['left'])
                section.right_margin = Inches(self.config.margins['right'])
            else:
                # Fallback a pulgadas
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Encabezado y pie
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

            # Tamaño de página
            if 'page' in self.config.__dict__:
                page_width = self.config.margins.get('page', {}).get('width', 8.5)
                page_height = self.config.margins.get('page', {}).get('height', 11)
                section.page_width = Inches(page_width)
                section.page_height = Inches(page_height)

        logger.debug("Documento configurado con márgenes APA 7")

    def build(self, elements: List[MarkdownElement],
              metadata: Optional[Dict[str, Any]] = None) -> Document:
        """
        Construir documento completo

        Args:
            elements: Lista de elementos parseados
            metadata: Metadatos del documento (título, autor, etc.)

        Returns:
            Documento DOCX completo

        Example:
            >>> from markdown_parser import MarkdownParser
            >>> from apa7_config import APAConfig
            >>> parser = MarkdownParser()
            >>> elements = parser.parse_file('tesis.md')
            >>> config = APAConfig.default_student()
            >>> builder = DocumentBuilder(config)
            >>> doc = builder.build(elements, {'title': 'Mi Tesis'})
        """
        try:
            logger.info(f"Construyendo documento con {len(elements)} elementos")

            # Agregar portada si está habilitada
            if self.config.cover_page.get('enabled', True) and metadata:
                self._add_cover_page(metadata)

            # Procesar cada elemento
            for i, element in enumerate(elements):
                if i % 100 == 0:
                    logger.debug(f"Procesando elemento {i}/{len(elements)}")

                self._process_element(element)

            logger.info("Documento construido exitosamente")
            return self.document

        except Exception as e:
            logger.error(f"Error construyendo documento: {e}")
            raise BuildError(f"Error construyendo documento: {e}")

    def _add_cover_page(self, metadata: Dict[str, Any]) -> None:
        """Agregar portada"""
        # Espacios para centrado vertical
        for _ in range(8):
            self.document.add_paragraph()

        # Título (negrita, centrado)
        title = metadata.get('title', '')
        if title:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.name = self.config.font['name']
            run.font.size = Pt(self.config.font['size'])

            self.document.add_paragraph()

        # Autor
        author = metadata.get('author', '')
        if author:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(author)
            run.font.name = self.config.font['name']
            run.font.size = Pt(self.config.font['size'])

        # Institución
        institution = metadata.get('institution', '')
        if institution:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(institution)
            run.font.name = self.config.font['name']
            run.font.size = Pt(self.config.font['size'])

        # Elementos adicionales según tipo
        if self.config.document_type == 'student':
            # Curso
            course = metadata.get('course', '')
            if course:
                self.document.add_paragraph()
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(course)
                run.font.name = self.config.font['name']
                run.font.size = Pt(self.config.font['size'])

            # Instructor
            instructor = metadata.get('instructor', '')
            if instructor:
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(instructor)
                run.font.name = self.config.font['name']
                run.font.size = Pt(self.config.font['size'])

            # Fecha
            date = metadata.get('date', '')
            if date:
                for _ in range(3):
                    self.document.add_paragraph()
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(date)
                run.font.name = self.config.font['name']
                run.font.size = Pt(self.config.font['size'])

        # Salto de página
        self.document.add_page_break()
        logger.debug("Portada agregada")

    def _process_element(self, element: MarkdownElement) -> None:
        """Procesar un elemento individual"""
        try:
            if element.type == ElementType.HEADING_1:
                self._add_heading(element, 'Heading 1')
            elif element.type == ElementType.HEADING_2:
                self._add_heading(element, 'Heading 2')
            elif element.type == ElementType.HEADING_3:
                self._add_heading(element, 'Heading 3')
            elif element.type == ElementType.HEADING_4:
                self._add_heading(element, 'Heading 4')
            elif element.type == ElementType.HEADING_5:
                self._add_heading(element, 'Heading 5')
            elif element.type == ElementType.PARAGRAPH:
                self._add_paragraph(element)
            elif element.type == ElementType.LIST_BULLET:
                self._add_bullet_list(element)
            elif element.type == ElementType.LIST_NUMBERED:
                self._add_numbered_list(element)
            elif element.type == ElementType.TABLE:
                self._add_table(element)
            elif element.type == ElementType.CODE_BLOCK:
                self._add_code_block(element)
            elif element.type == ElementType.BLOCKQUOTE:
                self._add_blockquote(element)
            elif element.type == ElementType.HORIZONTAL_RULE:
                # Ignorar separadores horizontales
                pass

        except Exception as e:
            logger.warning(f"Error procesando elemento en línea {element.line_number}: {e}")

    def _add_heading(self, element: MarkdownElement, style_name: str) -> None:
        """Agregar heading"""
        # Detectar si es sección de Referencias
        if 'referencia' in element.content.lower():
            self.in_references = True

        p = self.document.add_paragraph(element.content)
        p.style = style_name

    def _add_paragraph(self, element: MarkdownElement) -> None:
        """Agregar párrafo con formato inline"""
        p = self.document.add_paragraph()

        # Si estamos en referencias, usar estilo Reference
        if self.in_references:
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

        # Aplicar formato inline
        self._apply_inline_formatting(p, element.content)

    def _apply_inline_formatting(self, paragraph, text: str) -> None:
        """Aplicar formato inline (negrita, cursiva, código)"""
        parts = self.inline_formatter.parse(text)

        for part in parts:
            run = paragraph.add_run(part.text)

            if part.bold:
                run.bold = True
            if part.italic:
                run.italic = True
            if part.code:
                run.font.name = 'Courier New'
                run.font.size = Pt(11)

    def _add_bullet_list(self, element: MarkdownElement) -> None:
        """Agregar item de lista con bullet"""
        p = self.document.add_paragraph(style='List Bullet')
        self._apply_inline_formatting(p, element.content)

    def _add_numbered_list(self, element: MarkdownElement) -> None:
        """Agregar item de lista numerada"""
        p = self.document.add_paragraph(style='List Number')
        self._apply_inline_formatting(p, element.content)

    def _add_table(self, element: MarkdownElement) -> None:
        """Agregar tabla con formato APA"""
        rows_data = element.content

        if not rows_data or len(rows_data) < 1:
            return

        # Incrementar contador
        self.table_counter += 1

        # Número de tabla
        p_num = self.document.add_paragraph(f"Tabla {self.table_counter}")
        p_num.style = 'Table Number'

        # Crear tabla
        num_cols = len(rows_data[0])
        table = self.document.add_table(rows=len(rows_data), cols=num_cols)

        # Llenar tabla
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text

                    # Formato de celda
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in para.runs:
                            run.font.name = self.config.font['name']
                            run.font.size = Pt(self.config.font['size'])
                            if i == 0:  # Primera fila en negrita
                                run.bold = True

        # Espacio después
        self.document.add_paragraph()

    def _add_code_block(self, element: MarkdownElement) -> None:
        """Agregar bloque de código"""
        p = self.document.add_paragraph(element.content)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(0)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    def _add_blockquote(self, element: MarkdownElement) -> None:
        """Agregar blockquote (cita en bloque)"""
        p = self.document.add_paragraph(element.content)
        p.style = 'Quote'

    def save(self, output_path: str) -> None:
        """
        Guardar documento

        Args:
            output_path: Ruta donde guardar el documento
        """
        try:
            self.document.save(output_path)
            logger.info(f"Documento guardado: {output_path}")
        except Exception as e:
            logger.error(f"Error guardando documento: {e}")
            raise BuildError(f"Error guardando documento: {e}")
