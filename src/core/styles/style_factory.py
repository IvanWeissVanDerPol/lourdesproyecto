#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Style Factory - Factory para crear y aplicar estilos APA 7

Este módulo proporciona una interfaz para aplicar definiciones de estilos
APA 7 a documentos DOCX.
"""

from typing import Dict, List
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.styles.style import _ParagraphStyle
import logging

from .apa_styles import APAStyleDefinition, ALL_APA_STYLES


logger = logging.getLogger(__name__)


class StyleFactory:
    """
    Factory para crear y aplicar estilos APA 7 a documentos DOCX

    Esta clase toma definiciones de estilos (APAStyleDefinition) y las
    aplica a un documento, creando o modificando los estilos de Word
    correspondientes.

    Attributes:
        document: Documento DOCX al que aplicar los estilos
        created_styles: Diccionario de estilos ya creados
    """

    def __init__(self, document: Document):
        """
        Inicializar factory con un documento

        Args:
            document: Documento DOCX al que aplicar estilos
        """
        self.document = document
        self.created_styles: Dict[str, _ParagraphStyle] = {}

    def apply_style(self, style_def: APAStyleDefinition) -> _ParagraphStyle:
        """
        Aplicar una definición de estilo al documento

        Args:
            style_def: Definición del estilo a aplicar

        Returns:
            Estilo creado o modificado

        Raises:
            Exception: Si hay error aplicando el estilo

        Example:
            >>> from docx import Document
            >>> from apa_styles import APA_HEADING_1
            >>> doc = Document()
            >>> factory = StyleFactory(doc)
            >>> style = factory.apply_style(APA_HEADING_1)
            >>> style.name
            'Heading 1'
        """
        try:
            # Obtener o crear estilo
            style = self._get_or_create_style(style_def.name)

            # Aplicar formato de fuente
            style.font.name = style_def.font_name
            style.font.size = style_def.font_size
            style.font.bold = style_def.bold
            style.font.italic = style_def.italic

            # Aplicar formato de párrafo
            pf = style.paragraph_format
            pf.alignment = style_def.alignment
            pf.line_spacing = style_def.line_spacing
            pf.space_before = style_def.space_before
            pf.space_after = style_def.space_after
            pf.first_line_indent = style_def.first_line_indent
            pf.left_indent = style_def.left_indent
            pf.keep_with_next = style_def.keep_with_next

            # Guardar en cache
            self.created_styles[style_def.name] = style

            logger.debug(f"Estilo '{style_def.name}' aplicado correctamente")
            return style

        except Exception as e:
            logger.error(f"Error aplicando estilo '{style_def.name}': {e}")
            raise

    def apply_all_apa_styles(self) -> int:
        """
        Aplicar todos los estilos APA 7 al documento

        Returns:
            Número de estilos aplicados

        Example:
            >>> from docx import Document
            >>> doc = Document()
            >>> factory = StyleFactory(doc)
            >>> count = factory.apply_all_apa_styles()
            >>> print(f"Aplicados {count} estilos")
            Aplicados 11 estilos
        """
        count = 0
        for style_def in ALL_APA_STYLES:
            try:
                self.apply_style(style_def)
                count += 1
            except Exception as e:
                logger.warning(f"No se pudo aplicar estilo '{style_def.name}': {e}")
                continue

        logger.info(f"Aplicados {count}/{len(ALL_APA_STYLES)} estilos APA 7")
        return count

    def apply_styles_from_list(self, style_defs: List[APAStyleDefinition]) -> int:
        """
        Aplicar una lista específica de estilos

        Args:
            style_defs: Lista de definiciones de estilos a aplicar

        Returns:
            Número de estilos aplicados exitosamente

        Example:
            >>> from apa_styles import APA_HEADING_1, APA_HEADING_2
            >>> styles = [APA_HEADING_1, APA_HEADING_2]
            >>> factory.apply_styles_from_list(styles)
            2
        """
        count = 0
        for style_def in style_defs:
            try:
                self.apply_style(style_def)
                count += 1
            except Exception as e:
                logger.warning(f"No se pudo aplicar estilo '{style_def.name}': {e}")
                continue

        return count

    def _get_or_create_style(self, style_name: str) -> _ParagraphStyle:
        """
        Obtener o crear un estilo en el documento

        Args:
            style_name: Nombre del estilo

        Returns:
            Estilo existente o recién creado

        Raises:
            Exception: Si hay error creando el estilo
        """
        try:
            # Intentar obtener estilo existente
            return self.document.styles[style_name]
        except KeyError:
            # Crear nuevo estilo
            logger.debug(f"Creando nuevo estilo: {style_name}")
            return self.document.styles.add_style(
                style_name,
                WD_STYLE_TYPE.PARAGRAPH
            )

    def get_applied_styles(self) -> List[str]:
        """
        Obtener lista de nombres de estilos aplicados

        Returns:
            Lista de nombres de estilos

        Example:
            >>> factory.apply_all_apa_styles()
            >>> styles = factory.get_applied_styles()
            >>> 'Heading 1' in styles
            True
        """
        return list(self.created_styles.keys())

    def has_style(self, style_name: str) -> bool:
        """
        Verificar si un estilo ha sido aplicado

        Args:
            style_name: Nombre del estilo a verificar

        Returns:
            True si el estilo existe, False en caso contrario

        Example:
            >>> factory.has_style('Heading 1')
            False
            >>> factory.apply_style(APA_HEADING_1)
            >>> factory.has_style('Heading 1')
            True
        """
        return style_name in self.created_styles

    def reset(self) -> None:
        """
        Limpiar cache de estilos aplicados

        Nota: Esto no elimina los estilos del documento, solo limpia
        el cache interno del factory.
        """
        self.created_styles.clear()
        logger.debug("Cache de estilos limpiado")


class APAStyleValidator:
    """
    Validador de conformidad de estilos con APA 7

    Esta clase verifica que los estilos aplicados cumplan con los
    estándares APA 7.
    """

    @staticmethod
    def validate_style(style: _ParagraphStyle,
                       expected_def: APAStyleDefinition) -> bool:
        """
        Validar que un estilo cumpla con su definición esperada

        Args:
            style: Estilo de Word a validar
            expected_def: Definición esperada del estilo

        Returns:
            True si el estilo es conforme, False en caso contrario

        Example:
            >>> from docx import Document
            >>> doc = Document()
            >>> factory = StyleFactory(doc)
            >>> style = factory.apply_style(APA_HEADING_1)
            >>> validator = APAStyleValidator()
            >>> validator.validate_style(style, APA_HEADING_1)
            True
        """
        checks = [
            style.font.name == expected_def.font_name,
            style.font.size == expected_def.font_size,
            style.font.bold == expected_def.bold,
            style.font.italic == expected_def.italic,
            style.paragraph_format.alignment == expected_def.alignment,
        ]

        return all(checks)

    @staticmethod
    def validate_document_styles(document: Document) -> Dict[str, bool]:
        """
        Validar todos los estilos APA en un documento

        Args:
            document: Documento a validar

        Returns:
            Diccionario con resultados de validación para cada estilo

        Example:
            >>> from docx import Document
            >>> doc = Document()
            >>> factory = StyleFactory(doc)
            >>> factory.apply_all_apa_styles()
            >>> validator = APAStyleValidator()
            >>> results = validator.validate_document_styles(doc)
            >>> all(results.values())
            True
        """
        results = {}

        for style_def in ALL_APA_STYLES:
            try:
                style = document.styles[style_def.name]
                results[style_def.name] = APAStyleValidator.validate_style(
                    style, style_def
                )
            except KeyError:
                results[style_def.name] = False

        return results
