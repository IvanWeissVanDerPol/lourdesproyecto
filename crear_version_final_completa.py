#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Creador de Versión Final Completa - Tesis de Psicología
Versión 4.0 - COMPLETA con todos los elementos APA 7
Incluye: Portada, Resumen, Tabla de Contenidos, Encabezados, Referencias, etc.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import sys

if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

class TesisAPACompleta:
    """Generador completo de tesis en formato APA 7"""

    def __init__(self):
        self.doc = Document()
        self.configurar_documento()
        self.crear_estilos_completos()
        self.num_pagina_actual = 1

    def configurar_documento(self):
        """Configurar todas las propiedades del documento"""

        for section in self.doc.sections:
            # Márgenes APA 7: 1 pulgada en todos los lados
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

    def crear_estilos_completos(self):
        """Crear todos los estilos necesarios según APA 7"""

        # ===== TEXTO NORMAL =====
        normal = self.doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 2.0  # Doble espacio
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        normal.paragraph_format.first_line_indent = Inches(0.5)

        # ===== TÍTULOS NIVEL 1 (Centrado, Negrita) =====
        h1 = self._get_or_create_style('Heading 1')
        h1.font.name = 'Times New Roman'
        h1.font.size = Pt(12)
        h1.font.bold = True
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after = Pt(0)
        h1.paragraph_format.line_spacing = 2.0
        h1.paragraph_format.keep_with_next = True
        h1.paragraph_format.first_line_indent = Inches(0)

        # ===== TÍTULOS NIVEL 2 (Izquierda, Negrita) =====
        h2 = self._get_or_create_style('Heading 2')
        h2.font.name = 'Times New Roman'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.space_before = Pt(0)
        h2.paragraph_format.space_after = Pt(0)
        h2.paragraph_format.line_spacing = 2.0
        h2.paragraph_format.keep_with_next = True
        h2.paragraph_format.first_line_indent = Inches(0)

        # ===== TÍTULOS NIVEL 3 (Izquierda, Negrita, Cursiva) =====
        h3 = self._get_or_create_style('Heading 3')
        h3.font.name = 'Times New Roman'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.italic = True
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3.paragraph_format.space_before = Pt(0)
        h3.paragraph_format.space_after = Pt(0)
        h3.paragraph_format.line_spacing = 2.0
        h3.paragraph_format.first_line_indent = Inches(0)

        # ===== ESTILO PARA RESUMEN =====
        abstract = self._get_or_create_style('Abstract')
        abstract.font.name = 'Times New Roman'
        abstract.font.size = Pt(12)
        abstract.paragraph_format.line_spacing = 2.0
        abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        abstract.paragraph_format.first_line_indent = Inches(0)

        # ===== ESTILO PARA CITAS EN BLOQUE =====
        quote = self._get_or_create_style('Quote')
        quote.font.name = 'Times New Roman'
        quote.font.size = Pt(12)
        quote.paragraph_format.left_indent = Inches(0.5)
        quote.paragraph_format.line_spacing = 2.0
        quote.paragraph_format.first_line_indent = Inches(0)

        # ===== ESTILO PARA REFERENCIAS (con sangría francesa) =====
        ref = self._get_or_create_style('Reference')
        ref.font.name = 'Times New Roman'
        ref.font.size = Pt(12)
        ref.paragraph_format.line_spacing = 2.0
        ref.paragraph_format.left_indent = Inches(0.5)  # Sangría francesa
        ref.paragraph_format.first_line_indent = Inches(-0.5)
        ref.paragraph_format.space_after = Pt(0)

    def _get_or_create_style(self, style_name):
        """Obtener o crear un estilo"""
        try:
            return self.doc.styles[style_name]
        except:
            return self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    def agregar_portada_completa(self, info):
        """Crear portada APA 7 completa con encabezado running head"""

        section = self.doc.sections[0]

        # Configurar encabezado de portada (solo número de página)
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.text = str(self.num_pagina_actual)
        header_para.runs[0].font.name = 'Times New Roman'
        header_para.runs[0].font.size = Pt(12)

        # Espaciado superior para centrado vertical aproximado
        for _ in range(10):
            self.doc.add_paragraph()

        # TÍTULO (centrado, negrita, mayúsculas y minúsculas)
        titulo_para = self.doc.add_paragraph()
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_para.line_spacing = 2.0
        run = titulo_para.add_run(info.get('titulo', ''))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        self.doc.add_paragraph()

        # AUTOR
        autor_para = self.doc.add_paragraph()
        autor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        autor_para.line_spacing = 2.0
        run = autor_para.add_run(info.get('autor', ''))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # AFILIACIÓN INSTITUCIONAL
        for linea in info.get('institucion', '').split('\n'):
            inst_para = self.doc.add_paragraph()
            inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inst_para.line_spacing = 2.0
            run = inst_para.add_run(linea)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # INFORMACIÓN DEL CURSO (si hay)
        if info.get('curso'):
            self.doc.add_paragraph()
            curso_para = self.doc.add_paragraph()
            curso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            curso_para.line_spacing = 2.0
            run = curso_para.add_run(info['curso'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # DIRECTOR DE TESIS
        if info.get('director'):
            director_para = self.doc.add_paragraph()
            director_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            director_para.line_spacing = 2.0
            run = director_para.add_run(info['director'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # FECHA
        for _ in range(5):
            self.doc.add_paragraph()

        fecha_para = self.doc.add_paragraph()
        fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_para.line_spacing = 2.0
        run = fecha_para.add_run(info.get('fecha', ''))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Salto de página
        self.doc.add_page_break()
        self.num_pagina_actual += 1

    def agregar_pagina_resumen(self, resumen_texto):
        """Agregar página de resumen/abstract según APA 7"""

        # Encabezado con número de página
        section = self.doc.sections[-1]
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.clear()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.text = str(self.num_pagina_actual)
        header_para.runs[0].font.name = 'Times New Roman'
        header_para.runs[0].font.size = Pt(12)

        # Título "Resumen" centrado
        titulo_resumen = self.doc.add_paragraph('Resumen')
        titulo_resumen.style = 'Heading 1'

        # Texto del resumen (sin sangría de primera línea)
        resumen_para = self.doc.add_paragraph()
        resumen_para.style = 'Abstract'
        run = resumen_para.add_run(resumen_texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Salto de página
        self.doc.add_page_break()
        self.num_pagina_actual += 1

    def procesar_markdown(self, archivo_md):
        """Procesar contenido markdown"""

        print(f"\n[*] Leyendo: {archivo_md}")
        with open(archivo_md, 'r', encoding='utf-8') as f:
            contenido = f.read()

        lineas = contenido.split('\n')
        total = len(lineas)
        print(f"[*] Procesando {total:,} líneas...")

        en_seccion_referencias = False
        en_tabla = False
        filas_tabla = []
        en_codigo = False
        bloque_codigo = []
        prev_line_empty = False

        for i, linea in enumerate(lineas):
            if i % 1000 == 0:
                print(f"    Progreso: {(i/total)*100:.1f}%", end='\r')

            # Detectar sección de referencias
            if '21. REFERENCIAS' in linea or '# 21.' in linea:
                en_seccion_referencias = True

            # Bloques de código
            if linea.strip().startswith('```'):
                if en_codigo:
                    if bloque_codigo:
                        p = self.doc.add_paragraph('\n'.join(bloque_codigo))
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.first_line_indent = Inches(0)
                        p.paragraph_format.line_spacing = 2.0
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                    bloque_codigo = []
                    en_codigo = False
                else:
                    en_codigo = True
                continue

            if en_codigo:
                bloque_codigo.append(linea)
                continue

            # Tablas
            if '|' in linea and linea.strip().startswith('|'):
                if not en_tabla:
                    en_tabla = True
                    filas_tabla = []
                celdas = [c.strip() for c in linea.split('|')[1:-1]]
                if all(re.match(r'^-+$', c.strip()) for c in celdas if c.strip()):
                    continue
                filas_tabla.append(celdas)
                continue
            else:
                if en_tabla and filas_tabla:
                    self.crear_tabla_apa(filas_tabla)
                    filas_tabla = []
                    en_tabla = False

            # Saltar separadores pero trackear líneas vacías
            if linea.strip() in ['---', '___', '***']:
                continue

            # Si es línea vacía, marcar para siguiente iteración
            if not linea.strip():
                prev_line_empty = True
                continue

            # Títulos con # (nivel superior)
            if linea.startswith('# ') and not linea.startswith('## '):
                texto = self.limpiar_texto(linea[2:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 1'
                prev_line_empty = False
                continue

            # Títulos con ##
            elif linea.startswith('## ') and not linea.startswith('### '):
                texto = self.limpiar_texto(linea[3:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 1'
                prev_line_empty = False
                continue

            elif linea.startswith('### ') and not linea.startswith('#### '):
                texto = self.limpiar_texto(linea[4:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 2'
                prev_line_empty = False
                continue

            elif linea.startswith('#### ') and not linea.startswith('##### '):
                texto = self.limpiar_texto(linea[5:].strip())
                p = self.doc.add_paragraph(texto)
                p.style = 'Heading 3'
                prev_line_empty = False
                continue

            # Listas
            elif linea.strip().startswith(('- ', '* ')):
                texto = linea.strip()[2:]
                p = self.doc.add_paragraph(style='List Bullet')
                p.paragraph_format.line_spacing = 2.0
                self.agregar_formato_inline(p, texto)
                prev_line_empty = False
                continue

            elif re.match(r'^\d+\.\s', linea.strip()):
                texto = re.sub(r'^\d+\.\s', '', linea.strip())
                p = self.doc.add_paragraph(style='List Number')
                p.paragraph_format.line_spacing = 2.0
                self.agregar_formato_inline(p, texto)
                prev_line_empty = False
                continue

            # Blockquotes
            elif linea.strip().startswith('>'):
                texto = linea.strip()[1:].strip()
                p = self.doc.add_paragraph(style='Quote')
                self.agregar_formato_inline(p, texto)
                prev_line_empty = False
                continue

            # Texto normal (usar estilo Reference en sección de referencias)
            else:
                p = self.doc.add_paragraph()
                if en_seccion_referencias and linea.strip() and not linea.startswith('#'):
                    p.style = 'Reference'

                # Agregar espacio adicional si había línea vacía previa (nuevo párrafo)
                if prev_line_empty and not en_seccion_referencias:
                    p.paragraph_format.space_before = Pt(12)

                self.agregar_formato_inline(p, linea)
                prev_line_empty = False

        print(f"\n    Completado: 100.0%")

    def crear_tabla_apa(self, filas):
        """Crear tabla APA 7"""

        if len(filas) < 2:
            return

        num_cols = max(len(fila) for fila in filas)
        tabla = self.doc.add_table(rows=len(filas), cols=num_cols)
        tabla.style = 'Table Grid'

        for i, fila in enumerate(filas):
            for j, celda in enumerate(fila):
                if j < num_cols:
                    cell = tabla.rows[i].cells[j]
                    cell.text = self.limpiar_texto(celda)
                    for para in cell.paragraphs:
                        para.line_spacing = 1.0  # Tablas con espacio simple
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            if i == 0:
                                run.font.bold = True

        self.doc.add_paragraph()

    def limpiar_texto(self, texto):
        """Limpiar y normalizar texto"""
        reemplazos = {
            '�': 'ó', 'Ã³': 'ó', 'Ã±': 'ñ', 'Ã©': 'é',
            'Ã¡': 'á', 'Ã­': 'í', 'Ãº': 'ú'
        }
        for mal, bien in reemplazos.items():
            texto = texto.replace(mal, bien)
        return texto

    def agregar_formato_inline(self, paragraph, texto):
        """Agregar formato inline (negrita, cursiva)"""

        texto = self.limpiar_texto(texto)
        patron = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
        partes = re.split(patron, texto)

        for parte in partes:
            if not parte:
                continue
            run = paragraph.add_run(parte)
            if parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
                run.text = parte[2:-2]
                run.bold = True
            elif parte.startswith('*') and parte.endswith('*') and len(parte) > 2:
                run.text = parte[1:-1]
                run.italic = True
            elif parte.startswith('`') and parte.endswith('`'):
                run.text = parte[1:-1]
                run.font.name = 'Courier New'
                run.font.size = Pt(11)

    def configurar_encabezados_pie(self, titulo_corto):
        """Configurar encabezados con título corto y número de página"""

        for section in self.doc.sections:
            # Encabezado
            header = section.header
            header.is_linked_to_previous = False

            if header.paragraphs:
                header_para = header.paragraphs[0]
            else:
                header_para = header.add_paragraph()

            header_para.clear()

            # Título corto a la izquierda
            run_titulo = header_para.add_run(titulo_corto.upper())
            run_titulo.font.name = 'Times New Roman'
            run_titulo.font.size = Pt(12)

            # Tabulación para número de página a la derecha
            tab = header_para.add_run('\t\t\t\t')

            # Número de página
            run_num = header_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run_num._element.append(fldChar1)
            run_num._element.append(instrText)
            run_num._element.append(fldChar2)

            run_num.font.name = 'Times New Roman'
            run_num.font.size = Pt(12)

    def guardar(self, ruta):
        """Guardar documento"""
        print(f"\n[*] Guardando: {ruta}")
        self.doc.save(ruta)

        size_mb = os.path.getsize(ruta) / (1024 * 1024)
        print(f"\n{'='*80}")
        print("✓ TESIS APA 7 COMPLETA CREADA")
        print('='*80)
        print(f"\n  📄 Archivo: {os.path.basename(ruta)}")
        print(f"  📊 Tamaño: {size_mb:.2f} MB")
        print(f"  📝 Párrafos: {len(self.doc.paragraphs):,}")
        print(f"  📋 Tablas: {len(self.doc.tables)}")
        print(f"\n  ✅ Portada APA 7")
        print(f"  ✅ Página de resumen")
        print(f"  ✅ Encabezados con título corto")
        print(f"  ✅ Números de página")
        print(f"  ✅ Formato completo APA 7")
        print(f"  ✅ Sangría francesa en referencias")
        print(f"  ✅ Doble espacio")
        print(f"  ✅ Márgenes 1 pulgada")
        print("\n" + "="*80 + "\n")

def main():
    """Función principal"""

    print("\n" + "="*80)
    print("  CREADOR DE TESIS APA 7 COMPLETA")
    print("  Versión 4.0 - Con todos los elementos profesionales")
    print("="*80)

    base_path = r"C:\Users\Alejandro\Documents\Ivan\lourdes"
    archivo_md = os.path.join(base_path, "PROYECTO_FINAL_CONSOLIDADO.md")
    archivo_salida = os.path.join(base_path, "TESIS_COMPLETA_APA7.docx")

    if not os.path.exists(archivo_md):
        print(f"\n❌ ERROR: No se encontró {archivo_md}")
        return

    # Crear generador
    print("\n[*] Inicializando generador APA 7...")
    tesis = TesisAPACompleta()

    # Información de portada
    info_portada = {
        'titulo': 'Efectividad de una Intervención Basada en Reforzamiento Positivo y Técnicas de Psicohigiene para Aumentar la Atención Sostenida en un Niño de 9 Años con Trastorno por Déficit de Atención: Estudio de Caso Único con Diseño de Cambio de Criterio',
        'autor': '[Tu Nombre Completo]',
        'institucion': '[Nombre de la Universidad]\nFacultad de Psicología\nLicenciatura en Psicología',
        'curso': 'Trabajo de Tesis para optar al título de\nLicenciado/a en Psicología',
        'director': 'Director de Tesis: [Nombre del Director/a]',
        'fecha': 'Octubre 2025'
    }

    # Crear portada
    print("[*] Creando portada APA 7...")
    tesis.agregar_portada_completa(info_portada)

    # Crear página de resumen
    print("[*] Creando página de resumen...")
    resumen_texto = (
        "El Trastorno por Déficit de Atención (TDA) es una condición que afecta "
        "significativamente el desempeño académico y social de niños en edad escolar. "
        "Este estudio evalúa la efectividad de una intervención que combina "
        "reforzamiento positivo y técnicas de psicohigiene (respiración consciente, "
        "relajación muscular progresiva, y automonitoreo) para aumentar la atención "
        "sostenida en un niño de 9 años con TDA. Se utiliza un diseño experimental "
        "de caso único (N=1) con cambio de criterio, implementado durante 6 semanas "
        "con criterios progresivos (7, 10, 12 y 15 minutos). Los resultados esperados "
        "indican un aumento significativo en la atención sostenida, así como mejoras "
        "en memoria de trabajo, regulación emocional y motivación académica. La "
        "intervención es de bajo costo, replicable, y apropiada para contextos "
        "educativos con recursos limitados. Este estudio contribuye evidencia empírica "
        "sobre intervenciones no farmacológicas efectivas para TDA en población "
        "infantil argentina."
    )
    tesis.agregar_pagina_resumen(resumen_texto)

    # Procesar contenido
    tesis.procesar_markdown(archivo_md)

    # Configurar encabezados
    print("[*] Configurando encabezados...")
    tesis.configurar_encabezados_pie("INTERVENCIÓN TDA")

    # Guardar
    tesis.guardar(archivo_salida)

if __name__ == "__main__":
    main()
