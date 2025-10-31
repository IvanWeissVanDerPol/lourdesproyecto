#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertidor de Markdown a DOCX con formato profesional
Genera un documento Word completo del proyecto de investigación
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def crear_estilos(doc):
    """Crear estilos personalizados para el documento"""

    # Estilo para títulos nivel 1 (##)
    try:
        heading1 = doc.styles['Heading 1']
    except:
        heading1 = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

    heading1.font.name = 'Arial'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.keep_with_next = True

    # Estilo para títulos nivel 2 (###)
    try:
        heading2 = doc.styles['Heading 2']
    except:
        heading2 = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)

    heading2.font.name = 'Arial'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 51, 102)
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(10)
    heading2.paragraph_format.keep_with_next = True

    # Estilo para títulos nivel 3 (####)
    try:
        heading3 = doc.styles['Heading 3']
    except:
        heading3 = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)

    heading3.font.name = 'Arial'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(51, 51, 51)
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)

    # Estilo para texto normal
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def procesar_markdown(archivo_md, archivo_docx):
    """Convertir archivo markdown a DOCX con formato"""

    print("\n" + "="*70)
    print("  CREADOR DE DOCUMENTO WORD - PROYECTO DE INVESTIGACIÓN")
    print("="*70)
    print(f"\n[*] Leyendo: {archivo_md}")

    # Leer el archivo markdown
    with open(archivo_md, 'r', encoding='utf-8') as f:
        contenido = f.read()

    # Crear documento
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Crear estilos
    crear_estilos(doc)

    print("[*] Procesando contenido...")

    # Procesar línea por línea
    lineas = contenido.split('\n')
    total_lineas = len(lineas)

    en_tabla = False
    filas_tabla = []
    en_bloque_codigo = False
    bloque_codigo = []

    for i, linea in enumerate(lineas):
        # Mostrar progreso cada 1000 líneas
        if i % 1000 == 0:
            porcentaje = (i / total_lineas) * 100
            print(f"    Procesando: {porcentaje:.1f}% ({i}/{total_lineas} líneas)")

        # Detectar bloque de código
        if linea.strip().startswith('```'):
            if en_bloque_codigo:
                # Fin del bloque de código
                if bloque_codigo:
                    p = doc.add_paragraph('\n'.join(bloque_codigo))
                    p.style = 'Normal'
                    p.paragraph_format.left_indent = Inches(0.5)
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(51, 51, 51)
                bloque_codigo = []
                en_bloque_codigo = False
            else:
                # Inicio del bloque de código
                en_bloque_codigo = True
            continue

        if en_bloque_codigo:
            bloque_codigo.append(linea)
            continue

        # Detectar tablas markdown
        if '|' in linea and linea.strip().startswith('|'):
            if not en_tabla:
                en_tabla = True
                filas_tabla = []

            # Procesar fila de tabla
            celdas = [c.strip() for c in linea.split('|')[1:-1]]

            # Saltar línea separadora (|---|---|)
            if all(re.match(r'^-+$', c.strip()) for c in celdas if c.strip()):
                continue

            filas_tabla.append(celdas)
            continue
        else:
            # Si había una tabla, crearla ahora
            if en_tabla and filas_tabla:
                if len(filas_tabla) > 1:  # Al menos encabezado + 1 fila
                    num_cols = len(filas_tabla[0])
                    tabla = doc.add_table(rows=len(filas_tabla), cols=num_cols)
                    tabla.style = 'Light Grid Accent 1'

                    # Llenar tabla
                    for idx_fila, fila in enumerate(filas_tabla):
                        for idx_col, celda in enumerate(fila):
                            if idx_col < num_cols:
                                cell = tabla.rows[idx_fila].cells[idx_col]
                                cell.text = celda

                                # Formato para encabezado
                                if idx_fila == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True

                filas_tabla = []
                en_tabla = False

        # Saltar líneas vacías
        if not linea.strip():
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_paragraph()
            continue

        # Detectar títulos
        if linea.startswith('# '):
            # Título nivel 1 - muy raro, probablemente sea título de sección principal
            p = doc.add_paragraph(linea[2:].strip())
            p.style = 'Heading 1'
            continue

        elif linea.startswith('## '):
            # Título nivel 2
            p = doc.add_paragraph(linea[3:].strip())
            p.style = 'Heading 1'
            continue

        elif linea.startswith('### '):
            # Título nivel 3
            p = doc.add_paragraph(linea[4:].strip())
            p.style = 'Heading 2'
            continue

        elif linea.startswith('#### '):
            # Título nivel 4
            p = doc.add_paragraph(linea[5:].strip())
            p.style = 'Heading 3'
            continue

        # Detectar listas
        elif linea.strip().startswith('- ') or linea.strip().startswith('* '):
            texto = linea.strip()[2:]
            p = doc.add_paragraph(texto, style='List Bullet')
            continue

        elif re.match(r'^\d+\.\s', linea.strip()):
            texto = re.sub(r'^\d+\.\s', '', linea.strip())
            p = doc.add_paragraph(texto, style='List Number')
            continue

        # Detectar separadores (---)
        elif linea.strip() == '---' or linea.strip() == '___':
            # Agregar espacio
            doc.add_paragraph()
            continue

        # Detectar blockquotes (>)
        elif linea.strip().startswith('>'):
            texto = linea.strip()[1:].strip()
            p = doc.add_paragraph(texto)
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(102, 102, 102)
            continue

        # Texto normal
        else:
            p = doc.add_paragraph()

            # Procesar formato inline (negrita, cursiva, código)
            texto = linea

            # Reemplazar ** por negrita
            partes = re.split(r'(\*\*.*?\*\*)', texto)
            for parte in partes:
                if parte.startswith('**') and parte.endswith('**'):
                    run = p.add_run(parte[2:-2])
                    run.bold = True
                elif parte.startswith('*') and parte.endswith('*') and not parte.startswith('**'):
                    run = p.add_run(parte[1:-1])
                    run.italic = True
                elif parte.startswith('`') and parte.endswith('`'):
                    run = p.add_run(parte[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(199, 37, 78)
                else:
                    p.add_run(parte)

    print(f"    Procesando: 100.0% ({total_lineas}/{total_lineas} líneas)")

    # Guardar documento
    print(f"\n[*] Guardando documento: {archivo_docx}")
    doc.save(archivo_docx)

    # Obtener información del archivo
    size_bytes = os.path.getsize(archivo_docx)
    size_kb = size_bytes / 1024
    size_mb = size_kb / 1024

    num_paragraphs = len(doc.paragraphs)
    num_tables = len(doc.tables)

    print("\n" + "="*70)
    print("[OK] Documento creado exitosamente!")
    print("="*70)
    print(f"\n    Archivo:    {archivo_docx}")
    print(f"    Tamaño:     {size_mb:.2f} MB ({size_kb:.1f} KB)")
    print(f"    Párrafos:   {num_paragraphs:,}")
    print(f"    Tablas:     {num_tables}")
    print(f"    Páginas:    ~{num_paragraphs // 45} (estimado)")
    print("\n" + "="*70)
    print("\n[*] El documento está listo para:")
    print("    - Revisión y edición en Microsoft Word")
    print("    - Ajustes de formato final")
    print("    - Impresión")
    print("    - Entrega a profesores")
    print("\n")

if __name__ == "__main__":
    # Rutas de archivos
    base_path = r"C:\Users\Alejandro\Documents\Ivan\lourdes"
    archivo_md = os.path.join(base_path, "PROYECTO_FINAL_CONSOLIDADO.md")
    archivo_docx = os.path.join(base_path, "PROYECTO_FINAL_COMPLETO.docx")

    # Verificar que existe el archivo markdown
    if not os.path.exists(archivo_md):
        print(f"\n[ERROR] No se encontró el archivo: {archivo_md}")
        exit(1)

    # Procesar
    procesar_markdown(archivo_md, archivo_docx)
