"""
Optimize Table Column Widths Based on Content
Analyzes content and adjusts column widths to minimize word wrapping
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class TableWidthOptimizer:
    def __init__(self, input_path, output_path):
        self.doc = Document(input_path)
        self.output_path = output_path
        self.changes_made = []

    def optimize_all_tables(self):
        """Optimize all tables for full width and proper column distribution"""
        print("=" * 80)
        print("OPTIMIZING TABLE WIDTHS FOR CONTENT")
        print("=" * 80)

        self.analyze_and_optimize_tables()
        self.save_document()
        self.print_summary()

    def analyze_and_optimize_tables(self):
        """Analyze content and optimize each table"""
        print("\nAnalyzing and optimizing tables...")

        for table_num, table in enumerate(self.doc.tables, 1):
            print(f"\n  Table {table_num}:")

            num_cols = len(table.columns)
            num_rows = len(table.rows)

            # Calculate content-based weights for each column
            col_weights = self.calculate_column_weights(table)

            # Total available width: 6.5 inches (8.5" page - 2" margins)
            total_width_emu = Inches(6.5)

            # Distribute width based on content weights
            col_widths = []
            total_weight = sum(col_weights)

            for weight in col_weights:
                # Minimum width: 0.5 inches
                min_width = Inches(0.5)
                # Calculated width based on content
                calculated_width = int((weight / total_weight) * total_width_emu)
                # Use the larger of minimum or calculated
                final_width = max(min_width, calculated_width)
                col_widths.append(final_width)

            # Ensure total width equals available width
            actual_total = sum(col_widths)
            if actual_total != total_width_emu:
                # Adjust the largest column to make up the difference
                diff = total_width_emu - actual_total
                max_idx = col_weights.index(max(col_weights))
                col_widths[max_idx] = int(col_widths[max_idx] + diff)

            # Apply widths to columns
            for col_idx, col in enumerate(table.columns):
                col.width = col_widths[col_idx]

            # Print column distribution
            col_info = []
            for i, width in enumerate(col_widths):
                inches = width / 914400  # Convert EMU to inches
                weight = col_weights[i]
                col_info.append(f"Col{i+1}={inches:.1f}\" (weight={weight})")

            print(f"    {num_rows}×{num_cols} - Widths: {', '.join(col_info)}")

            self.changes_made.append(f"Table {table_num}: Optimized {num_cols} columns")

    def calculate_column_weights(self, table):
        """
        Calculate relative weight for each column based on content length
        Returns list of weights (higher = needs more space)
        """
        num_cols = len(table.columns)
        col_max_lengths = [0] * num_cols

        # Sample rows to determine content length
        # Check header + up to 10 data rows
        rows_to_check = min(11, len(table.rows))

        for row_idx in range(rows_to_check):
            row = table.rows[row_idx]

            for col_idx, cell in enumerate(row.cells):
                # Get text content from all paragraphs in cell
                cell_text = ' '.join([p.text for p in cell.paragraphs])
                cell_length = len(cell_text)

                # Track maximum length for this column
                if cell_length > col_max_lengths[col_idx]:
                    col_max_lengths[col_idx] = cell_length

        # Convert lengths to weights
        # Use square root to reduce extreme differences
        col_weights = []
        for length in col_max_lengths:
            if length == 0:
                weight = 10  # Minimum weight for empty columns
            else:
                # Square root scaling + minimum base weight
                weight = max(10, int((length ** 0.6) * 2))
            col_weights.append(weight)

        # Special handling for common patterns
        col_weights = self.adjust_weights_for_patterns(table, col_weights)

        return col_weights

    def adjust_weights_for_patterns(self, table, weights):
        """Adjust weights based on common table patterns"""

        # Check if first column is a narrow index/number column
        if len(table.columns) > 1:
            first_col_samples = []
            for row_idx in range(min(5, len(table.rows))):
                cell_text = table.rows[row_idx].cells[0].text.strip()
                first_col_samples.append(cell_text)

            # If first column looks like numbers/short labels, reduce weight
            if all(len(text) < 15 for text in first_col_samples):
                weights[0] = int(weights[0] * 0.6)

        # Check if last column is numeric (often narrow)
        if len(table.columns) > 1:
            last_col_idx = len(table.columns) - 1
            last_col_samples = []

            for row_idx in range(min(5, len(table.rows))):
                cell_text = table.rows[row_idx].cells[last_col_idx].text.strip()
                last_col_samples.append(cell_text)

            # If last column looks numeric/short, reduce weight
            numeric_count = sum(1 for text in last_col_samples
                              if text.replace('.', '').replace(',', '').replace('%', '').replace('-', '').isdigit())

            if numeric_count >= 3 or all(len(text) < 10 for text in last_col_samples):
                weights[last_col_idx] = int(weights[last_col_idx] * 0.7)

        return weights

    def save_document(self):
        """Save the optimized document"""
        print("\n\nSaving optimized document...")
        self.doc.save(self.output_path)
        print(f"✓ Document saved: {self.output_path}")

    def print_summary(self):
        """Print summary of changes"""
        print("\n" + "=" * 80)
        print("TABLE WIDTH OPTIMIZATION COMPLETE")
        print("=" * 80)

        print(f"\nTables optimized: {len(self.doc.tables)}")
        print(f"\nKey improvements:")
        print("  ✓ Tables use full available width (6.5 inches)")
        print("  ✓ Column widths adjusted based on content length")
        print("  ✓ Word wrapping minimized")
        print("  ✓ Narrow columns (numbers, indices) optimized")
        print("  ✓ Wide columns (text, descriptions) get more space")

        print("\n" + "=" * 80)
        print("OPTIMIZATION DETAILS")
        print("=" * 80)
        print("\nWidth Distribution Strategy:")
        print("  • Content analysis: Measures actual text length in each column")
        print("  • Smart weighting: Square root scaling to avoid extreme differences")
        print("  • Pattern detection: Identifies numeric/index columns (narrower)")
        print("  • Full width usage: Distributes all 6.5\" across columns")
        print("  • Minimum width: 0.5\" per column guaranteed")

        print("\n" + "=" * 80)
        print("Next: Open document and verify tables look optimal")
        print("=" * 80)

if __name__ == "__main__":
    input_file = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7_FINAL.docx"
    output_file = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7_FINAL_OPTIMIZED.docx"

    print("Starting table width optimization...")
    print(f"Input: {input_file}")
    print(f"Output: {output_file}\n")

    optimizer = TableWidthOptimizer(input_file, output_file)
    optimizer.optimize_all_tables()

    print("\n" + "=" * 80)
    print("DONE!")
    print("=" * 80)
    print("\nYour tables now:")
    print("  • Use full page width (6.5 inches)")
    print("  • Have content-optimized column widths")
    print("  • Minimize word wrapping")
    print("\nRecommendation:")
    print("  1. Open the optimized document")
    print("  2. Review tables for appearance")
    print("  3. If satisfied, replace the original")
    print("=" * 80)
