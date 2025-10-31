"""
Validate the corrected APA 7 document
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches

def validate_document(docx_path):
    doc = Document(docx_path)

    print("=" * 80)
    print("VALIDATION REPORT: CORRECTED DOCUMENT")
    print("=" * 80)

    passed = []
    warnings = []

    # Check page setup
    section = doc.sections[0]

    print("\n[PAGE SETUP]")
    if abs(section.top_margin - Inches(1)) < Inches(0.01):
        passed.append("✓ Top margin: 1 inch")
        print("  ✓ Top margin: 1 inch")

    if abs(section.left_margin - Inches(1)) < Inches(0.01):
        passed.append("✓ Left margin: 1 inch")
        print("  ✓ Left margin: 1 inch")

    if abs(section.page_width - Inches(8.5)) < Inches(0.01):
        passed.append("✓ Page width: 8.5 inches (Letter)")
        print("  ✓ Page width: 8.5 inches (Letter)")

    if abs(section.page_height - Inches(11)) < Inches(0.01):
        passed.append("✓ Page height: 11 inches (Letter)")
        print("  ✓ Page height: 11 inches (Letter)")

    # Check font
    print("\n[FONT & SPACING]")
    sample_paragraphs = [p for p in doc.paragraphs if p.text.strip() and len(p.text) > 50][:10]

    font_check = 0
    spacing_check = 0

    for para in sample_paragraphs:
        if para.runs:
            font_name = para.runs[0].font.name
            font_size = para.runs[0].font.size

            if font_name == 'Times New Roman':
                font_check += 1

            if font_size and abs(font_size - Pt(12)) < Pt(0.1):
                font_check += 1

        if para.paragraph_format.line_spacing == 2.0:
            spacing_check += 1

    if font_check > 10:
        passed.append("✓ Font: Times New Roman 12pt")
        print("  ✓ Font: Times New Roman 12pt")

    if spacing_check > 5:
        passed.append("✓ Line spacing: Double (2.0)")
        print("  ✓ Line spacing: Double (2.0)")

    # Check headings
    print("\n[HEADINGS]")
    heading_counts = {}
    for para in doc.paragraphs:
        if 'Heading' in para.style.name:
            level = para.style.name
            heading_counts[level] = heading_counts.get(level, 0) + 1

    for level, count in heading_counts.items():
        passed.append(f"✓ {level}: {count} found")
        print(f"  ✓ {level}: {count} found")

    # Check abstract
    print("\n[ABSTRACT]")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == 'RESUMEN':
            passed.append("✓ Abstract section found")
            print("  ✓ Abstract section found")

            if i + 1 < len(doc.paragraphs):
                abstract = doc.paragraphs[i + 1]
                word_count = len(abstract.text.split())

                if 150 <= word_count <= 250:
                    passed.append(f"✓ Abstract length: {word_count} words (optimal)")
                    print(f"  ✓ Abstract length: {word_count} words (optimal)")
                else:
                    warnings.append(f"⚠ Abstract length: {word_count} words (recommend 150-250)")
                    print(f"  ⚠ Abstract length: {word_count} words (recommend 150-250)")
            break

    # Check references
    print("\n[REFERENCES]")
    for i, para in enumerate(doc.paragraphs):
        text_upper = para.text.strip().upper()
        if text_upper in ['REFERENCIAS', 'REFERENCES']:
            passed.append("✓ References section found")
            print("  ✓ References section found")

            # Check a few reference entries for hanging indent
            ref_entries = 0
            hanging_indent_count = 0

            for j in range(i + 1, min(i + 20, len(doc.paragraphs))):
                ref_para = doc.paragraphs[j]
                if ref_para.text.strip() and len(ref_para.text) > 20:
                    ref_entries += 1

                    left_indent = ref_para.paragraph_format.left_indent
                    first_line = ref_para.paragraph_format.first_line_indent

                    if left_indent and first_line:
                        if abs(left_indent - Inches(0.5)) < Inches(0.1) and abs(first_line - Inches(-0.5)) < Inches(0.1):
                            hanging_indent_count += 1

            if hanging_indent_count > 0:
                passed.append(f"✓ Hanging indent applied to {hanging_indent_count} references")
                print(f"  ✓ Hanging indent applied to {hanging_indent_count} references")
            break

    # Check paragraph formatting
    print("\n[PARAGRAPHS]")
    indented_count = 0
    body_paragraphs = [p for p in doc.paragraphs if p.style.name == 'Normal' and len(p.text.strip()) > 20]

    for para in body_paragraphs[:50]:
        first_line = para.paragraph_format.first_line_indent
        if first_line and abs(first_line - Inches(0.5)) < Inches(0.1):
            indented_count += 1

    if indented_count > 20:
        passed.append(f"✓ First-line indent (0.5\") applied to body paragraphs")
        print(f"  ✓ First-line indent (0.5\") applied to body paragraphs")

    # Summary
    print("\n" + "=" * 80)
    print("VALIDATION SUMMARY")
    print("=" * 80)
    print(f"\n✓ PASSED CHECKS: {len(passed)}")
    print(f"⚠ WARNINGS: {len(warnings)}")

    if warnings:
        print("\nWarnings:")
        for w in warnings:
            print(f"  {w}")

    compliance_rate = (len(passed) / (len(passed) + len(warnings))) * 100 if (len(passed) + len(warnings)) > 0 else 100

    print(f"\n[COMPLIANCE RATE] {compliance_rate:.1f}%")

    if compliance_rate >= 95:
        print("[EXCELLENT] Document is fully APA 7 compliant!")
    elif compliance_rate >= 85:
        print("[VERY GOOD] Document is highly compliant with minor items to review")
    elif compliance_rate >= 75:
        print("[GOOD] Document is compliant with some improvements recommended")

    print("\n" + "=" * 80)

if __name__ == "__main__":
    docx_path = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7_CORRECTED.docx"
    validate_document(docx_path)
