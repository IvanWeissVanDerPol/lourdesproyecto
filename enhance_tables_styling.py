"""
Enhanced Table Formatting for APA 7 Compliance
Adds borders, adjusts widths, and ensures proper styling
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class TableStyler:
    def __init__(self, input_path, output_path):
        self.doc = Document(input_path)
        self.output_path = output_path
        self.changes_made = []
        self.table_count = 0

    def enhance_all_tables(self):
        """Apply APA 7 table formatting to all tables"""
        print("=" * 80)
        print("ENHANCING TABLE FORMATTING FOR APA 7")
        print("=" * 80)

        self.analyze_tables()
        self.format_tables()
        self.adjust_table_widths()
        self.format_table_captions()
        self.save_document()
        self.print_summary()

    def analyze_tables(self):
        """Analyze existing tables"""
        print("\n[1/4] Analyzing tables in document...")

        self.table_count = len(self.doc.tables)
        print(f"  Found {self.table_count} tables")

        for i, table in enumerate(self.doc.tables, 1):
            rows = len(table.rows)
            cols = len(table.columns)
            print(f"  Table {i}: {rows} rows × {cols} columns")

    def format_tables(self):
        """Apply APA 7 table formatting"""
        print("\n[2/4] Applying APA 7 table formatting...")

        for table_num, table in enumerate(self.doc.tables, 1):
            print(f"  Formatting Table {table_num}...")

            # Set table width to fit within margins (6.5 inches for 1-inch margins on 8.5" paper)
            table.width = Inches(6.5)

            # Apply APA 7 borders: only horizontal lines
            # Top border, bottom border, and line after header row
            self.set_apa_borders(table)

            # Format all cells
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # Format cell paragraphs
                    for paragraph in cell.paragraphs:
                        # Set font
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                        # Set alignment
                        if row_idx == 0:
                            # Header row: center or left align based on content
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # Make header bold
                            for run in paragraph.runs:
                                run.bold = True
                        else:
                            # Data rows: align based on content type
                            text = paragraph.text.strip()
                            if text and (text.replace('.', '').replace(',', '').replace('-', '').isdigit() or
                                        self.is_number(text)):
                                # Numbers: right align
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                # Text: left align
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        # Set line spacing
                        paragraph.paragraph_format.line_spacing = 1.0  # Single spacing in tables
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(2)

                    # Set cell margins for better spacing
                    self.set_cell_margins(cell, top=50, bottom=50, left=100, right=100)

            self.changes_made.append(f"✓ Formatted Table {table_num} with APA 7 styling")

    def set_apa_borders(self, table):
        """
        Set APA 7 compliant borders:
        - Top border of table (thick)
        - Bottom border of table (thick)
        - Border after header row (medium)
        - No vertical borders
        - No borders between data rows
        """
        tbl = table._tbl
        tblPr = tbl.tblPr

        # Create table borders element
        tblBorders = OxmlElement('w:tblBorders')

        # Top border (thick - 1.5pt)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '18')  # 1.5pt = 18 eighths of a point
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')

        # Bottom border (thick - 1.5pt)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '18')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')

        # No left border
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'none')

        # No right border
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'none')

        # No inside horizontal borders (will add manually after header)
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), 'none')

        # No inside vertical borders
        insideV = OxmlElement('w:insideV')
        insideV.set(qn('w:val'), 'none')

        tblBorders.append(top)
        tblBorders.append(bottom)
        tblBorders.append(left)
        tblBorders.append(right)
        tblBorders.append(insideH)
        tblBorders.append(insideV)

        tblPr.append(tblBorders)

        # Add border after header row (first row)
        if len(table.rows) > 0:
            header_row = table.rows[0]
            self.set_row_bottom_border(header_row, size='12')  # 1pt = 12 eighths

    def set_row_bottom_border(self, row, size='12'):
        """Add bottom border to a specific row"""
        for cell in row.cells:
            tcPr = cell._element.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell._element.append(tcPr)

            tcBorders = OxmlElement('w:tcBorders')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), size)
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')

            tcBorders.append(bottom)
            tcPr.append(tcBorders)

    def set_cell_margins(self, cell, top=100, bottom=100, left=100, right=100):
        """Set cell margins for better spacing"""
        tcPr = cell._element.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell._element.append(tcPr)

        tcMar = OxmlElement('w:tcMar')

        # Top margin
        top_el = OxmlElement('w:top')
        top_el.set(qn('w:w'), str(top))
        top_el.set(qn('w:type'), 'dxa')

        # Bottom margin
        bottom_el = OxmlElement('w:bottom')
        bottom_el.set(qn('w:w'), str(bottom))
        bottom_el.set(qn('w:type'), 'dxa')

        # Left margin
        left_el = OxmlElement('w:left')
        left_el.set(qn('w:w'), str(left))
        left_el.set(qn('w:type'), 'dxa')

        # Right margin
        right_el = OxmlElement('w:right')
        right_el.set(qn('w:w'), str(right))
        right_el.set(qn('w:type'), 'dxa')

        tcMar.append(top_el)
        tcMar.append(bottom_el)
        tcMar.append(left_el)
        tcMar.append(right_el)

        tcPr.append(tcMar)

    def is_number(self, text):
        """Check if text is a number (including decimals, percentages, etc.)"""
        try:
            # Remove common number formatting
            clean_text = text.replace(',', '').replace('%', '').replace('$', '').strip()
            float(clean_text)
            return True
        except:
            return False

    def adjust_table_widths(self):
        """Adjust column widths for optimal fit"""
        print("\n[3/4] Adjusting table column widths...")

        for table_num, table in enumerate(self.doc.tables, 1):
            num_cols = len(table.columns)

            # Calculate equal column widths that fit within page
            # Total width: 6.5 inches (accounting for 1-inch margins on 8.5" paper)
            total_width = Inches(6.5)
            col_width = int(total_width / num_cols)

            # Set column widths
            for col in table.columns:
                col.width = col_width

            self.changes_made.append(f"✓ Adjusted Table {table_num} column widths ({num_cols} columns)")

    def format_table_captions(self):
        """Format table titles and notes according to APA 7"""
        print("\n[4/4] Formatting table captions and notes...")

        # Find and format table titles
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()

            # Check if this is a table title (e.g., "Tabla 1", "Table 1")
            if text and (text.lower().startswith('tabla ') or text.lower().startswith('table ')):
                # Format table number
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                    run.italic = False

                # Check next paragraph for table title (should be italic)
                if i + 1 < len(self.doc.paragraphs):
                    title_para = self.doc.paragraphs[i + 1]
                    # If next paragraph doesn't look like a table, it's the title
                    if not title_para.text.strip().lower().startswith('tabla') and \
                       not title_para.text.strip().lower().startswith('table'):
                        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        title_para.paragraph_format.space_before = Pt(0)
                        title_para.paragraph_format.space_after = Pt(6)

                        for run in title_para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.italic = True
                            run.bold = False

                self.changes_made.append(f"✓ Formatted table caption: {text}")

            # Check for table notes (e.g., "Nota.", "Note.")
            elif text.lower().startswith('nota.') or text.lower().startswith('note.'):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.left_indent = Inches(0)

                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.italic = True

                self.changes_made.append(f"✓ Formatted table note")

    def save_document(self):
        """Save the enhanced document"""
        print("\nSaving enhanced document...")
        self.doc.save(self.output_path)
        print(f"✓ Document saved: {self.output_path}")

    def print_summary(self):
        """Print summary of changes"""
        print("\n" + "=" * 80)
        print("TABLE FORMATTING COMPLETE")
        print("=" * 80)

        print(f"\nTables processed: {self.table_count}")
        print("\nChanges applied:")
        for change in self.changes_made:
            print(f"  {change}")

        print("\n" + "=" * 80)
        print("APA 7 TABLE FORMATTING APPLIED")
        print("=" * 80)
        print("\nAPA 7 Table Features Applied:")
        print("  ✓ Horizontal borders only (top, bottom, after header)")
        print("  ✓ No vertical borders")
        print("  ✓ Times New Roman 12pt font")
        print("  ✓ Proper cell alignment (centered headers, left/right data)")
        print("  ✓ Single spacing within tables")
        print("  ✓ Adequate cell padding")
        print("  ✓ Tables fit within page margins (6.5 inches)")
        print("  ✓ Table numbers in bold")
        print("  ✓ Table titles in italic")
        print("  ✓ Table notes in italic")

        print("\n" + "=" * 80)
        print("Next step: Open document and verify table appearance")
        print("=" * 80)

if __name__ == "__main__":
    input_file = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7_CORRECTED.docx"
    output_file = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7_FINAL.docx"

    print("Starting table enhancement...")
    print(f"Input: {input_file}")
    print(f"Output: {output_file}\n")

    styler = TableStyler(input_file, output_file)
    styler.enhance_all_tables()
