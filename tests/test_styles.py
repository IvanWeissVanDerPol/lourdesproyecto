#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tests for APA Styles Module

Tests unitarios para el módulo de estilos APA 7.
"""

import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.core.styles.apa_styles import (
    APAStyleDefinition,
    APA_NORMAL,
    APA_HEADING_1,
    APA_HEADING_2,
    APA_HEADING_3,
    APA_HEADING_4,
    APA_HEADING_5,
    APA_QUOTE,
    APA_REFERENCE,
    get_style_by_name,
    get_heading_style
)
from src.core.styles.style_factory import StyleFactory


class TestAPAStyleDefinitions:
    """Tests para definiciones de estilos APA"""

    def test_apa_normal_properties(self):
        """Test: APA_NORMAL tiene propiedades correctas"""
        assert APA_NORMAL.name == 'Normal'
        assert APA_NORMAL.font_name == 'Times New Roman'
        assert APA_NORMAL.font_size == Pt(12)
        assert APA_NORMAL.line_spacing == 2.0
        assert APA_NORMAL.first_line_indent == Inches(0.5)

    def test_heading_1_is_bold_centered(self):
        """Test: Heading 1 es negrita y centrado"""
        assert APA_HEADING_1.bold == True
        assert APA_HEADING_1.italic == False
        assert APA_HEADING_1.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_heading_2_is_bold_left(self):
        """Test: Heading 2 es negrita e izquierda"""
        assert APA_HEADING_2.bold == True
        assert APA_HEADING_2.italic == False
        assert APA_HEADING_2.alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_heading_3_is_bold_italic(self):
        """Test: Heading 3 es negrita y cursiva"""
        assert APA_HEADING_3.bold == True
        assert APA_HEADING_3.italic == True

    def test_heading_4_has_indent(self):
        """Test: Heading 4 tiene sangría"""
        assert APA_HEADING_4.left_indent == Inches(0.5)
        assert APA_HEADING_4.bold == True
        assert APA_HEADING_4.italic == True

    def test_heading_5_has_indent(self):
        """Test: Heading 5 tiene sangría"""
        assert APA_HEADING_5.left_indent == Inches(0.5)
        assert APA_HEADING_5.bold == True
        assert APA_HEADING_5.italic == True

    def test_reference_has_hanging_indent(self):
        """Test: Reference tiene sangría francesa"""
        assert APA_REFERENCE.first_line_indent == Inches(-0.5)
        assert APA_REFERENCE.left_indent == Inches(0.5)

    def test_quote_has_indent(self):
        """Test: Quote tiene sangría"""
        assert APA_QUOTE.left_indent == Inches(0.5)
        assert APA_QUOTE.first_line_indent == Inches(0)

    def test_get_style_by_name(self):
        """Test: get_style_by_name encuentra estilos"""
        style = get_style_by_name('Heading 1')
        assert style is not None
        assert style.name == 'Heading 1'

        style = get_style_by_name('NonExistent')
        assert style is None

    def test_get_heading_style(self):
        """Test: get_heading_style devuelve estilos correctos"""
        assert get_heading_style(1) == APA_HEADING_1
        assert get_heading_style(2) == APA_HEADING_2
        assert get_heading_style(3) == APA_HEADING_3
        assert get_heading_style(4) == APA_HEADING_4
        assert get_heading_style(5) == APA_HEADING_5
        assert get_heading_style(6) is None


class TestStyleFactory:
    """Tests para StyleFactory"""

    def test_apply_style_creates_style(self):
        """Test: apply_style crea estilo en documento"""
        doc = Document()
        factory = StyleFactory(doc)

        style = factory.apply_style(APA_HEADING_1)

        assert style.name == 'Heading 1'
        assert style.font.bold == True

    def test_apply_all_apa_styles(self):
        """Test: apply_all_apa_styles crea todos los estilos"""
        doc = Document()
        factory = StyleFactory(doc)

        count = factory.apply_all_apa_styles()

        assert count >= 9  # Al menos 9 estilos básicos
        assert factory.has_style('Normal')
        assert factory.has_style('Heading 1')
        assert factory.has_style('Reference')

    def test_style_factory_idempotent(self):
        """Test: aplicar estilo dos veces no causa error"""
        doc = Document()
        factory = StyleFactory(doc)

        factory.apply_style(APA_HEADING_1)
        factory.apply_style(APA_HEADING_1)  # No debe fallar

        assert factory.has_style('Heading 1')
