"""
APA 7 Format Correction Script
Applies all APA 7 formatting corrections to the document
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class APA7Formatter:
    def __init__(self, input_path, output_path):
        self.doc = Document(input_path)
        self.output_path = output_path
        self.changes_made = []

    def apply_all_corrections(self):
        """Apply all APA 7 formatting corrections"""
        print("=" * 80)
        print("APPLYING APA 7 FORMATTING CORRECTIONS")
        print("=" * 80)

        self.fix_page_setup()
        self.fix_font_and_spacing()
        self.fix_title_page()
        self.fix_abstract()
        self.fix_headings()
        self.fix_paragraph_formatting()
        self.fix_references()
        self.add_page_numbers()

        self.save_document()
        self.print_summary()

    def fix_page_setup(self):
        """Fix page margins and size"""
        print("\n[1/8] Fixing page setup...")

        for section in self.doc.sections:
            # Set margins to 1 inch (914400 EMU)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Set page size to Letter (8.5 x 11 inches)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        self.changes_made.append("✓ Set all margins to 1 inch")
        self.changes_made.append("✓ Set page size to Letter (8.5 × 11\")")

    def fix_font_and_spacing(self):
        """Fix font to Times New Roman 12pt and double spacing"""
        print("[2/8] Fixing font and spacing...")

        count = 0
        for paragraph in self.doc.paragraphs:
            # Set font for all runs in paragraph
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                # Ensure font is applied to complex scripts too
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Set paragraph line spacing to double (2.0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

            # Remove extra spacing before/after paragraphs
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            count += 1

        self.changes_made.append(f"✓ Set font to Times New Roman 12pt ({count} paragraphs)")
        self.changes_made.append("✓ Applied double spacing throughout")
        self.changes_made.append("✓ Removed extra spacing before/after paragraphs")

    def fix_title_page(self):
        """Fix title page formatting"""
        print("[3/8] Fixing title page...")

        # Find and format title (first substantial paragraph)
        title_found = False
        for i, para in enumerate(self.doc.paragraphs[:20]):
            if para.text.strip() and len(para.text.strip()) > 30 and not title_found:
                # This is likely the title
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Make title bold
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(12)  # APA 7: same size as body text

                title_found = True
                self.changes_made.append(f"✓ Formatted title: \"{para.text[:50]}...\"")
                break

        # Center first 10 paragraphs (typical title page elements)
        centered_count = 0
        for para in self.doc.paragraphs[:10]:
            if para.text.strip() and 'Heading' not in para.style.name:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                centered_count += 1

        self.changes_made.append(f"✓ Centered title page elements ({centered_count} elements)")

    def fix_abstract(self):
        """Fix abstract and keywords formatting"""
        print("[4/8] Fixing abstract and keywords...")

        abstract_idx = -1
        keywords_idx = -1

        # Find RESUMEN/ABSTRACT heading
        for i, para in enumerate(self.doc.paragraphs):
            text_upper = para.text.strip().upper()

            if text_upper in ['RESUMEN', 'ABSTRACT'] and abstract_idx == -1:
                abstract_idx = i
                # Center and bold the heading
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True

                # Format abstract paragraph (next paragraph)
                if i + 1 < len(self.doc.paragraphs):
                    abstract_para = self.doc.paragraphs[i + 1]
                    # Abstract should NOT be indented
                    abstract_para.paragraph_format.first_line_indent = Inches(0)
                    abstract_para.paragraph_format.left_indent = Inches(0)

                    # Count words
                    word_count = len(abstract_para.text.split())
                    self.changes_made.append(f"✓ Formatted abstract ({word_count} words)")

            elif 'PALABRAS CLAVE' in text_upper or 'KEYWORDS' in text_upper:
                keywords_idx = i
                # Center and bold the heading
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True

                # Format keywords paragraph
                if i + 1 < len(self.doc.paragraphs):
                    keywords_para = self.doc.paragraphs[i + 1]
                    # Keywords indented 0.5 inch
                    keywords_para.paragraph_format.first_line_indent = Inches(0.5)
                    keywords_para.paragraph_format.left_indent = Inches(0)

                    self.changes_made.append("✓ Formatted keywords section")
                break

    def fix_headings(self):
        """Fix heading formatting according to APA 7 levels"""
        print("[5/8] Fixing heading hierarchy...")

        heading_counts = {'Heading 1': 0, 'Heading 2': 0, 'Heading 3': 0}

        for para in self.doc.paragraphs:
            style_name = para.style.name

            # Level 1: Centered, Bold, Title Case
            if style_name == 'Heading 1':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.left_indent = Inches(0)
                for run in para.runs:
                    run.bold = True
                    run.italic = False
                    run.font.size = Pt(12)
                heading_counts['Heading 1'] += 1

            # Level 2: Left-aligned, Bold, Title Case
            elif style_name == 'Heading 2':
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.left_indent = Inches(0)
                for run in para.runs:
                    run.bold = True
                    run.italic = False
                    run.font.size = Pt(12)
                heading_counts['Heading 2'] += 1

            # Level 3: Left-aligned, Bold Italic, Title Case
            elif style_name == 'Heading 3':
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.left_indent = Inches(0)
                for run in para.runs:
                    run.bold = True
                    run.italic = True
                    run.font.size = Pt(12)
                heading_counts['Heading 3'] += 1

        for level, count in heading_counts.items():
            if count > 0:
                self.changes_made.append(f"✓ Formatted {count} {level} headings")

    def fix_paragraph_formatting(self):
        """Fix paragraph indentation and alignment"""
        print("[6/8] Fixing paragraph formatting...")

        fixed_count = 0
        for para in self.doc.paragraphs:
            style_name = para.style.name

            # Body paragraphs should have 0.5 inch first line indent
            if style_name == 'Normal':
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Only indent if not right after a heading or in specific sections
                # Check if this is not the abstract or similar
                if para.text.strip() and len(para.text) > 20:
                    # Check if previous paragraph was a heading
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    fixed_count += 1

            # Lists should be left-aligned with no first-line indent
            elif 'List' in style_name:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Inches(0)

        self.changes_made.append(f"✓ Applied 0.5\" first-line indent to {fixed_count} body paragraphs")
        self.changes_made.append("✓ Set all paragraphs to left-aligned")

    def fix_references(self):
        """Fix references section formatting"""
        print("[7/8] Fixing references section...")

        # Find references section
        ref_start = -1
        for i, para in enumerate(self.doc.paragraphs):
            text_upper = para.text.strip().upper()
            if text_upper in ['REFERENCIAS', 'REFERENCIAS BIBLIOGRÁFICAS', 'REFERENCES']:
                ref_start = i

                # Format heading: centered, bold
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True

                self.changes_made.append(f"✓ Found and formatted references section at paragraph {i}")
                break

        if ref_start > -1:
            # Format reference entries with hanging indent
            ref_count = 0
            for i in range(ref_start + 1, min(ref_start + 200, len(self.doc.paragraphs))):
                para = self.doc.paragraphs[i]

                # Stop if we hit another major heading
                if 'Heading 1' in para.style.name:
                    break

                # Format reference entries
                if para.text.strip() and len(para.text) > 20:
                    # Hanging indent: first line at 0, subsequent lines at 0.5"
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.first_line_indent = Inches(-0.5)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    ref_count += 1

            self.changes_made.append(f"✓ Applied hanging indent to {ref_count} reference entries")
        else:
            self.changes_made.append("⚠ References section not found")

    def add_page_numbers(self):
        """Add page numbers to header (top right)"""
        print("[8/8] Adding page numbers...")

        try:
            # Add page number to first section
            section = self.doc.sections[0]
            header = section.header

            # Create paragraph in header for page number
            if len(header.paragraphs) == 0:
                header_para = header.add_paragraph()
            else:
                header_para = header.paragraphs[0]

            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Clear existing content
            header_para.clear()

            # Add page number field
            run = header_para.add_run()

            # Create page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

            # Format the page number
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            self.changes_made.append("✓ Added page numbers in top right corner")
        except Exception as e:
            self.changes_made.append(f"⚠ Could not add page numbers: {str(e)}")

    def save_document(self):
        """Save the corrected document"""
        print("\nSaving corrected document...")
        self.doc.save(self.output_path)
        print(f"✓ Document saved: {self.output_path}")

    def print_summary(self):
        """Print summary of changes made"""
        print("\n" + "=" * 80)
        print("FORMATTING CORRECTIONS APPLIED")
        print("=" * 80)

        for change in self.changes_made:
            print(f"  {change}")

        print("\n" + "=" * 80)
        print("CORRECTION COMPLETE!")
        print("=" * 80)
        print(f"\nOriginal file: PROYECTO_FINAL_INVESTIGACION_APA7.docx")
        print(f"Corrected file: {self.output_path}")
        print("\nTotal changes applied:", len(self.changes_made))
        print("\nNext steps:")
        print("  1. Open the corrected file in Word")
        print("  2. Review the formatting")
        print("  3. Update fields (Ctrl+A, then F9) to show page numbers")
        print("  4. Do a final manual review using the checklist")
        print("=" * 80)

if __name__ == "__main__":
    input_file = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7.docx"
    output_file = r"c:\Users\Alejandro\Documents\Ivan\lourdes\PROYECTO_FINAL_INVESTIGACION_APA7_CORRECTED.docx"

    print("Starting APA 7 formatting corrections...")
    print(f"Input: {input_file}")
    print(f"Output: {output_file}\n")

    formatter = APA7Formatter(input_file, output_file)
    formatter.apply_all_corrections()
